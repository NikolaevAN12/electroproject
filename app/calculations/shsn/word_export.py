"""Формирование отчёта Word: ТКЗ 0,4 кВ, чувствительность АВ, карты селективности."""

from __future__ import annotations

import math
from typing import Any

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .calculator import _Calculator
from .diagrams import create_replacement_diagram, create_selectivity_map
from .models import (
    CB_CURVE_MULTIPLIERS,
    DISPLAY_COLUMN_RENAME,
    EXCLUDE_FROM_SOURCE_TABLE,
    NetworkElement,
    ScPointResult,
    ShsnInput,
    ShsnResult,
)
from .network import NetworkGraph

_REPORT_FONT = "Times New Roman"
_REPORT_FONT_SIZE_PT = 14


def _set_run_font(run: Any, *, size_pt: int = _REPORT_FONT_SIZE_PT) -> None:
    run.font.name = _REPORT_FONT
    run.font.size = Pt(size_pt)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for tag in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        r_fonts.set(qn(tag), _REPORT_FONT)


def _apply_paragraph_font(paragraph: Any) -> None:
    for run in paragraph.runs:
        _set_run_font(run)


def _apply_report_font(doc: Any) -> None:
    for name in ("Normal", "List Bullet", *(f"Heading {i}" for i in range(10))):
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        style.font.name = _REPORT_FONT
        style.font.size = Pt(_REPORT_FONT_SIZE_PT)

    for paragraph in doc.paragraphs:
        _apply_paragraph_font(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _apply_paragraph_font(paragraph)


def populate_shsn_document(doc: Any, input_model: ShsnInput, result_model: ShsnResult) -> None:
    graph = NetworkGraph(input_model.elements)
    calc = _Calculator(graph, k_temp=input_model.k_temp)
    _EnhancedReportGenerator(input_model.elements, result_model.sc_results, calc).generate(doc)


class _EnhancedReportGenerator:
    def __init__(
        self,
        elements: list[NetworkElement],
        results: dict[str, ScPointResult],
        calculator: _Calculator,
    ) -> None:
        self.elements = elements
        self.results = results
        self.calc = calculator
        self.doc = None
        self.fig_counter = 1
        self.table_counters: dict[int, int] = {}
        self.curves = CB_CURVE_MULTIPLIERS

    def generate(self, doc: Any) -> None:
        self.doc = doc
        title_text = "Расчет ТКЗ в сети 0,4 кВ и проверка чувствительности АВ"
        title = self._add_black_heading(title_text, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self._add_black_heading("Раздел 1: Исходные данные", level=1)
        self._add_table_caption(1, "Параметры элементов сети")
        display_cols = [
            c for c in DISPLAY_COLUMN_RENAME if c not in EXCLUDE_FROM_SOURCE_TABLE
        ]
        table = self.doc.add_table(rows=1, cols=len(display_cols))
        table.style = "Table Grid"
        for i, col in enumerate(display_cols):
            table.rows[0].cells[i].text = DISPLAY_COLUMN_RENAME[col]
        for el in self.elements:
            row_cells = table.add_row().cells
            values = {
                "Name": el.name,
                "Type": el.type,
                "Parent_Name": el.parent_name,
                "Phase_Type": el.phase_type,
                "U_nom": el.u_nom,
                "R1": el.r1,
                "X1": el.x1,
                "R0": el.r0,
                "X0": el.x0,
            }
            for i, col in enumerate(display_cols):
                val = values[col]
                row_cells[i].text = str(val) if not (isinstance(val, float) and math.isnan(val)) else ""

        self.doc.add_page_break()
        self._add_black_heading("Раздел 2: Схемы замещения", level=1)
        for name, res in self.results.items():
            self._add_black_heading(f"Точка КЗ: {name}", level=2)
            self.doc.add_picture(create_replacement_diagram(name, res, "1"), width=Inches(6.0))
            p = self.doc.add_paragraph(
                f"Рис. {self.fig_counter} Схема замещения прямой последовательности для {name}"
            )
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self.fig_counter += 1
            self.doc.add_picture(create_replacement_diagram(name, res, "0"), width=Inches(6.0))
            p = self.doc.add_paragraph(
                f"Рис. {self.fig_counter} Схема замещения нулевой последовательности для {name}"
            )
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self.fig_counter += 1

        self.doc.add_page_break()
        self._add_black_heading("Раздел 3: Ход расчета", level=1)
        for name, res in self.results.items():
            self._add_black_heading(f"Точка КЗ: {name}", level=2)
            self._write_calculation_steps(res)

        self.doc.add_page_break()
        self._add_black_heading("Раздел 4: Итоговая таблица результатов", level=1)
        self._add_table_caption(4, "Сводные данные по токам короткого замыкания")
        summ = self.doc.add_table(rows=1, cols=5)
        summ.style = "Table Grid"
        for i, h in enumerate(
            ["Точка КЗ", "U ном (кВ)", "Iкз(3) (кА)", "Iкз(1) холод (кА)", "Iкз(1) нагр (кА)"]
        ):
            summ.rows[0].cells[i].text = h
        for name, res in self.results.items():
            r = summ.add_row().cells
            r[0].text = str(name)
            r[1].text = f"{res.u_nom}"
            r[2].text = f"{res.i_sc_3:.3f}" if "3PH" in str(res.phase_type).upper() else "—"
            r[3].text = f"{res.i_sc_1_cold:.3f}"
            r[4].text = f"{res.i_sc_1_hot:.3f}"

        self.doc.add_page_break()
        self._add_black_heading("Раздел 5: Проверка чувствительности АВ", level=1)
        self._write_sensitivity_table()

        self.doc.add_page_break()
        self._add_black_heading("Раздел 6: Карты селективности АВ", level=1)
        self._write_selectivity_maps()
        _apply_report_font(self.doc)

    def _write_calculation_steps(self, res: ScPointResult) -> None:
        def get_sqrt3():
            return self._create_rad("", [self._create_t("3")])

        self.doc.add_paragraph("1. Определение суммарных сопротивлений (Ом):")
        for code, label in (("1", "прямая"), ("0", "нулевая")):
            self.doc.add_paragraph(f"Последовательность {label}:")
            r_attr = f"r{code}"
            x_attr = f"x{code}"
            r_sum_str = " + ".join([f"{getattr(e, r_attr):.4f}" for e in res.elements_data])
            x_sum_str = " + ".join([f"{getattr(e, x_attr):.4f}" for e in res.elements_data])
            r_sum = res.r_sum_1 if code == "1" else res.r_sum_0
            x_sum = res.x_sum_1 if code == "1" else res.x_sum_0
            self.doc.add_paragraph(f"ΣR{code} = {r_sum_str} = {r_sum:.4f} Ом")
            self.doc.add_paragraph(f"ΣX{code} = {x_sum_str} = {x_sum:.4f} Ом")
        self.doc.add_paragraph("2. Расчет тока КЗ:")
        if "3PH" in str(res.phase_type).upper():
            self.doc.add_paragraph("Трехфазное КЗ:", style="List Bullet")
            p1 = self.doc.add_paragraph()
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            num1 = [self._create_sub("U", "н")]
            den1 = [
                get_sqrt3(),
                self._create_t("·"),
                self._create_rad(
                    "",
                    [
                        self._create_sup(self._create_sub("ΣR", "1"), "2"),
                        self._create_t("+"),
                        self._create_sup(self._create_sub("ΣX", "1"), "2"),
                    ],
                ),
            ]
            self._add_math_structured(p1, [self._create_sub("I", "кз3"), self._create_t(" = "), self._create_frac(num1, den1)])
            p1c = self.doc.add_paragraph()
            p1c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            num1c = [self._create_t(str(res.u_nom))]
            den1c = [
                get_sqrt3(),
                self._create_t("·"),
                self._create_rad(
                    "",
                    [
                        self._create_sup(f"({res.r_sum_1:.4f})", "2"),
                        self._create_t("+"),
                        self._create_sup(f"({res.x_sum_1:.4f})", "2"),
                    ],
                ),
            ]
            self._add_math_structured(
                p1c,
                [
                    self._create_sub("I", "кз3"),
                    self._create_t(" = "),
                    self._create_frac(num1c, den1c),
                    self._create_t(f" = {res.i_sc_3:.3f} кА"),
                ],
            )
        self.doc.add_paragraph("Однофазное КЗ (без учета нагрева):", style="List Bullet")
        self.doc.add_paragraph(f"ΣR_петли = 2·ΣR1 + ΣR0 = {res.r_total_1ph_cold:.4f} Ом")
        self.doc.add_paragraph(f"ΣX_петли = 2·ΣX1 + ΣX0 = {res.x_total_1ph:.4f} Ом")
        p_c_s = self.doc.add_paragraph()
        p_c_s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        num_c_s = [self._create_sub("U", "н"), self._create_t("·"), get_sqrt3()]
        den_c_s = [
            self._create_rad(
                "",
                [
                    self._create_sup(self._create_sub("R", "петли"), "2"),
                    self._create_t("+"),
                    self._create_sup(self._create_sub("X", "петли"), "2"),
                ],
            )
        ]
        self._add_math_structured(
            p_c_s, [self._create_sub("I", "кз1_холод"), self._create_t(" = "), self._create_frac(num_c_s, den_c_s)]
        )
        p_c_v = self.doc.add_paragraph()
        p_c_v.alignment = WD_ALIGN_PARAGRAPH.CENTER
        num_c_v = [self._create_t(f"{res.u_nom}·"), get_sqrt3()]
        den_c_v = [
            self._create_rad(
                "",
                [
                    self._create_sup(f"({res.r_total_1ph_cold:.4f})", "2"),
                    self._create_t("+"),
                    self._create_sup(f"({res.x_total_1ph:.4f})", "2"),
                ],
            )
        ]
        self._add_math_structured(
            p_c_v,
            [
                self._create_sub("I", "кз1_холод"),
                self._create_t(" = "),
                self._create_frac(num_c_v, den_c_v),
                self._create_t(f" = {res.i_sc_1_cold:.3f} кА"),
            ],
        )
        self.doc.add_paragraph("Однофазное КЗ (с учетом нагрева):", style="List Bullet")
        p2s = self.doc.add_paragraph()
        p2s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        num2s = [self._create_sub("U", "н"), self._create_t("·"), get_sqrt3()]
        den2s = [
            self._create_rad(
                "",
                [
                    self._create_sup(self._create_sub("R", "петли_нагр"), "2"),
                    self._create_t("+"),
                    self._create_sup(self._create_sub("X", "петли"), "2"),
                ],
            )
        ]
        self._add_math_structured(
            p2s, [self._create_sub("I", "кз1_нагр"), self._create_t(" = "), self._create_frac(num2s, den2s)]
        )
        p2c = self.doc.add_paragraph()
        p2c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        num2c = [self._create_t(f"{res.u_nom}·"), get_sqrt3()]
        den2c = [
            self._create_rad(
                "",
                [
                    self._create_sup(f"({res.r_total_1ph_hot:.4f})", "2"),
                    self._create_t("+"),
                    self._create_sup(f"({res.x_total_1ph:.4f})", "2"),
                ],
            )
        ]
        self._add_math_structured(
            p2c,
            [
                self._create_sub("I", "кз1_нагр"),
                self._create_t(" = "),
                self._create_frac(num2c, den2c),
                self._create_t(f" = {res.i_sc_1_hot:.3f} кА"),
            ],
        )

    def _write_sensitivity_table(self) -> None:
        self._add_table_caption(5, "Результаты проверки чувствительности защитных аппаратов")
        sens_table = self.doc.add_table(rows=1, cols=11)
        sens_table.style = "Table Grid"
        headers = [
            "Присоединение",
            "Iном (А)",
            "Кратн.",
            "Хар-ка откл.",
            "Iкз3 нач (кА)",
            "Iкз1 кон.х (кА)",
            "Iкз1 кон.н (кА)",
            "Kч.х",
            "Kч.н",
            "Время (с)",
            "Результат",
        ]
        for i, h in enumerate(headers):
            sens_table.rows[0].cells[i].text = h

        automats = [el for el in self.elements if "автомат" in el.type.lower()]
        for el in automats:
            name = el.name
            node = el.as_graph_node()
            path_to_start = self.calc.graph.get_path_to_source(name)
            sc_start = self.calc.get_sc_params(path_to_start)
            children = [e for e in self.elements if e.parent_name == name]
            if children:
                sc_end = self.calc.get_sc_params(self.calc.graph.get_path_to_source(children[0].name))
            else:
                sc_end = sc_start

            i_nom = el.cb_nominal
            k_mult_input = el.cb_multiplier
            t_input = el.cb_time
            curve_input = el.cb_curve.upper().strip()

            actual_k_mult = (
                k_mult_input
                if (not math.isnan(k_mult_input) and k_mult_input > 0)
                else self.curves.get(curve_input, 10)
            )
            actual_t_cb = t_input if (not math.isnan(t_input) and t_input > 0) else 0.015

            if curve_input and curve_input != "NAN":
                char_text = curve_input
            elif (
                not math.isnan(i_nom)
                and not math.isnan(k_mult_input)
                and not math.isnan(t_input)
            ):
                char_text = "регул."
            else:
                char_text = "—"

            denom = (i_nom * actual_k_mult) / 1000.0 if not math.isnan(i_nom) else 0.0
            k_sens_cold = sc_end["i_sc_1_cold"] / denom if denom > 0 else 0.0
            k_sens_hot = sc_end["i_sc_1_hot"] / denom if denom > 0 else 0.0

            r_cells = sens_table.add_row().cells
            r_cells[0].text = str(name)
            r_cells[1].text = f"{i_nom}" if not math.isnan(i_nom) else ""
            r_cells[2].text = f"{actual_k_mult}"
            r_cells[3].text = char_text
            phase_type = str(el.phase_type).upper()
            r_cells[4].text = f"{sc_start['i_sc_3']:.3f}" if "3PH" in phase_type else "—"
            r_cells[5].text = f"{sc_end['i_sc_1_cold']:.3f}"
            r_cells[6].text = f"{sc_end['i_sc_1_hot']:.3f}"
            r_cells[7].text = f"{k_sens_cold:.2f}"
            r_cells[8].text = f"{k_sens_hot:.2f}"
            r_cells[9].text = f"{actual_t_cb}"
            r_cells[10].text = "Удовл." if k_sens_hot >= 1.5 else "Не удовл."

    def _write_selectivity_maps(self) -> None:
        processed_pairs: set[tuple[str, str]] = set()
        fig_idx = 1
        for _name, res in self.results.items():
            cb_path = [
                n
                for n in res.path
                if n in self.calc.graph.nodes
                and "автомат" in str(self.calc.graph.nodes[n]["Type"]).lower()
            ]
            for i in range(len(cb_path) - 1):
                up_name, dw_name = cb_path[i], cb_path[i + 1]
                if (up_name, dw_name) in processed_pairs:
                    continue
                processed_pairs.add((up_name, dw_name))
                up_node = self.calc.graph.nodes[up_name]
                dw_node = self.calc.graph.nodes[dw_name]

                path_to_dw_output = self.calc.graph.get_path_to_source(dw_name, include_start_node=True)
                sc_at_dw_output = self.calc.get_sc_params(path_to_dw_output)
                dw_phase = str(dw_node.get("Phase_Type", "3PH")).upper()
                i_max_val = (
                    sc_at_dw_output["i_sc_3"]
                    if "3PH" in dw_phase
                    else sc_at_dw_output["i_sc_1_cold"]
                )

                next_pt, is_next_cb = self.calc.graph.get_end_of_protected_zone(dw_name)
                path_to_next = self.calc.graph.get_path_to_source(
                    next_pt, include_start_node=not is_next_cb
                )
                sc_at_next_pt = self.calc.get_sc_params(path_to_next)
                i_min_val = sc_at_next_pt["i_sc_1_hot"]

                img, overlap = create_selectivity_map(up_node, dw_node, i_min_val, i_max_val)
                self.doc.add_picture(img, width=Inches(5.0))
                p_cap = self.doc.add_paragraph(
                    f"Рис. 6.{fig_idx} — Карта селективности пары {up_name} и {dw_name}"
                )
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                fig_idx += 1
                status = "СЕЛЕКТИВНОСТЬ НЕ СОБЛЮДАЕТСЯ" if overlap else "Селективность обеспечена"
                p_res = self.doc.add_paragraph(f"Вывод: {status}")
                if p_res.runs:
                    p_res.runs[0].bold = overlap

    def _get_table_num(self, section_num: int) -> str:
        self.table_counters[section_num] = self.table_counters.get(section_num, 0) + 1
        return f"{section_num}.{self.table_counters[section_num]}"

    def _add_table_caption(self, section_num: int, title: str) -> Any:
        num = self._get_table_num(section_num)
        p = self.doc.add_paragraph(f"Таблица {num} — {title}")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return p

    def _add_black_heading(self, text: str, level: int) -> Any:
        h = self.doc.add_heading(text, level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
        return h

    def _create_element(self, name: str) -> Any:
        return OxmlElement(f"m:{name}")

    def _create_t(self, text: str) -> Any:
        r = self._create_element("r")
        t = self._create_element("t")
        t.text = text
        r.append(t)
        return r

    def _create_sub(self, base_txt: str, sub_txt: str) -> Any:
        s_sub = self._create_element("sSub")
        e, sub = self._create_element("e"), self._create_element("sub")
        if isinstance(base_txt, str):
            e.append(self._create_t(base_txt))
        else:
            e.append(base_txt)
        sub.append(self._create_t(sub_txt))
        s_sub.append(e)
        s_sub.append(sub)
        return s_sub

    def _create_sup(self, base_node: Any, sup_txt: str) -> Any:
        s_sup = self._create_element("sSup")
        e, sup = self._create_element("e"), self._create_element("sup")
        if isinstance(base_node, str):
            e.append(self._create_t(base_node))
        else:
            e.append(base_node)
        sup.append(self._create_t(sup_txt))
        s_sup.append(e)
        s_sup.append(sup)
        return s_sup

    def _create_frac(self, num_nodes: list[Any], den_nodes: list[Any]) -> Any:
        f = self._create_element("f")
        num, den = self._create_element("num"), self._create_element("den")
        for n in num_nodes:
            num.append(n)
        for n in den_nodes:
            den.append(n)
        f.append(num)
        f.append(den)
        return f

    def _create_rad(self, deg_txt: str, deg_nodes: list[Any]) -> Any:
        rad = self._create_element("rad")
        deg, e = self._create_element("deg"), self._create_element("e")
        if deg_txt:
            deg.append(self._create_t(deg_txt))
        for n in deg_nodes:
            e.append(n)
        rad.append(deg)
        rad.append(e)
        return rad

    def _add_math_structured(self, paragraph: Any, nodes: list[Any]) -> None:
        omath_para = self._create_element("oMathPara")
        omath = self._create_element("oMath")
        for n in nodes:
            omath.append(n)
        omath_para.append(omath)
        paragraph._p.append(omath_para)
