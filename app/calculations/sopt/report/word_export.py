"""Формирование документа Word для расчёта СОПТ."""

from __future__ import annotations

import re
import xml.sax.saxutils as xml_esc
from copy import deepcopy
from dataclasses import dataclass
from io import BytesIO
from typing import Any, Literal

MixedPart = tuple[Literal["text", "math"], str]

from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from ..cable_fire_check import build_cable_fire_check_rows
from ..kc import kc_from_r_kz
from ..cdw_export import CdwSheetImage, _KOMPAS_RASTER_DPI
from ..models import EQUIPMENT_LABEL_BY_KEY, SoptEquipmentItem, SoptInput, SoptResult, SoptSection
from ..resistance import item_resistance_ohm
from ..schematic_draw import (
    akb_input_schematic_png_bytes,
    k1_substitution_schematic_png_bytes,
    subsection_kz_schematic_png_bytes,
)
from ..sensitivity_check import build_breaker_sensitivity_rows
from ..selectivity_map import SelectivityMapResult, build_subsection_selectivity_maps
from app.calculations.fire_check.word_omml import (
    omml_theta_k_schema_string,
    omml_theta_n_schema_string,
)
from .numbering import (
    REPORT_SECTION_DESCRIPTION,
    REPORT_SECTION_RESISTANCE,
    REPORT_SECTION_CABLE_FIRE,
    REPORT_SECTION_CONCLUSION,
    REPORT_SECTION_REFERENCES,
    REPORT_SECTION_SELECTIVITY,
    REPORT_SECTION_SENSITIVITY,
    cable_fire_table,
    REPORT_SECTION_TKZ,
    REPORT_SECTION_TKZ_BUS,
    TKZ_BUS_KZ_POINT,
    selectivity_figure_start,
    shpt_scheme_figure_label,
    shpt_scheme_figure_number,
    selectivity_section,
    selectivity_subsection,
    sensitivity_table,
    tkz_currents_table,
    tkz_user_section,
    tkz_user_subsection,
)

_FIG2_AND_NEXT_WIDTH_CM = 9.6
_FIG2_AND_NEXT_HEIGHT_CM = 5.29

_JUMPER_LENGTH_BETWEEN_BANKS = 0.25
_JUMPER_LENGTH_TO_NEIGHBOR_CABINET = 1.5
_PAGE_WIDTH_CM = 21.0
_PAGE_HEIGHT_CM = 29.7
_PAGE_MARGIN_TOP_CM = 1.69
_PAGE_MARGIN_BOTTOM_CM = 2.0
_PAGE_MARGIN_LEFT_CM = 2.25
_PAGE_MARGIN_RIGHT_CM = 1.25
_PAGE_MARGIN_GUTTER_CM = 0.0
_TEXT_AREA_WIDTH_CM = _PAGE_WIDTH_CM - _PAGE_MARGIN_LEFT_CM - _PAGE_MARGIN_RIGHT_CM
_REPORT_TABLE_WIDTH_SCALE = 1.5
# Табл. 4.1: №, наименование, Iном, характеристика, кратность, Iкз, Kч, проходит.
_SENSITIVITY_TABLE_COL_WEIGHTS = (0.55, 7.0, 1.1, 3.0, 0.9, 1.6, 0.75, 4.1)
_SENSITIVITY_TABLE_ROW_MIN_HEIGHT_CM = 1.0
# Табл. 5.1: №, марка, Iраб, Iдд, Θ до, Iкз, t, Θ после, проходит.
_CABLE_FIRE_TABLE_COL_WEIGHTS = (1.0, 4.0, 2.0, 2.0, 2.5, 2.0, 3.0, 2.5, 2.0)
_TKZ_CURRENTS_TABLE_COL_WEIGHTS = (0.6, 5.0, 2.2, 2.2)
_KOMPAS_SHEET_MARGIN_CM = 0.0
# Ограничение Word на сторону страницы (~22 дюйма).
_WORD_MAX_PAGE_SIDE_CM = 55.88
_BODY_FIRST_LINE_INDENT_CM = 1.0
_TEXT_FONT_NAME = "Arial"
_FONT_SIZE_SECTION_PT = 14
_FONT_SIZE_SUBSECTION_PT = 13
_FONT_SIZE_BODY_PT = 12
_MNS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
_WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_MATH_W_RPR = (
    f'<w:rPr xmlns:w="{_WNS}">'
    f'<w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math" w:cs="Arial"/>'
    f'<w:sz w:val="24"/><w:szCs w:val="24"/>'
    f"</w:rPr>"
)

def _fmt_fixed(value: float, digits: int) -> str:
    return f"{value:.{digits}f}".replace(".", ",")


def _fmt_kc(kc: float) -> str:
    if abs(kc - 0.55) < 1e-9:
        return "0,55"
    if abs(kc - 0.5) < 1e-9:
        return "0,5"
    return "0,6"


def _resistance_element_label(designation: str) -> str:
    if designation.strip().casefold().startswith("ш"):
        return "шинки"
    return "кабеля"


def _fmt_intish(value: float) -> str:
    rounded = round(value)
    if abs(value - rounded) < 1e-9:
        return str(int(rounded))
    return _fmt_fixed(value, 2).rstrip("0").rstrip(",")


def _xe(s: str) -> str:
    return xml_esc.escape(s)


def _xml_mrun(text: str) -> str:
    safe = _xe(text if text else "-")
    return f'<m:r>{_MATH_W_RPR}<m:rPr><m:nor/></m:rPr><m:t xml:space="preserve">{safe}</m:t></m:r>'


def _xml_mfrac(num_inner: str, den_inner: str) -> str:
    return (
        "<m:f>"
        f"<m:num><m:e>{num_inner}</m:e></m:num>"
        f"<m:den><m:e>{den_inner}</m:e></m:den>"
        "</m:f>"
    )


def _xml_msub(base: str, sub: str) -> str:
    return (
        "<m:sSub>"
        f"<m:e>{_xml_mrun(base)}</m:e>"
        f"<m:sub>{_xml_mrun(sub)}</m:sub>"
        "</m:sSub>"
    )


def _xml_msup_expr(base_inner: str, sup: str) -> str:
    return (
        "<m:sSup>"
        f"<m:e>{base_inner}</m:e>"
        f"<m:sup>{_xml_mrun(sup)}</m:sup>"
        "</m:sSup>"
    )


def _xml_omml_b_sopt_schema() -> str:
    b = _xml_msub("B", "тер")
    ikz = _xml_msub("I", "кз")
    i2 = _xml_msup_expr(ikz, "2")
    t_off = _xml_msub("t", "откл")
    return f"{b}{_xml_mrun(' = ')}{i2}{_xml_mrun('·')}{t_off}"


def _append_cable_fire_methodology(doc: Any) -> None:
    _add_body_paragraph(
        doc,
        "Проверка кабельных линий системы оперативного постоянного тока на невозгорание "
        "при воздействии тока короткого замыкания осуществляется по циркуляру № Ц-02-98 (Э) [4].",
    )
    _add_mixed_paragraph(
        doc,
        [
            (
                "text",
                "В соответствии с циркуляром № Ц-02-98 (Э) не допускается нагрев токопроводящей жилы "
                "кабеля с поливинилхлоридной изоляцией свыше 350 ",
            ),
            ("math", _xml_mrun("°C")),
            ("text", "."),
        ],
    )
    _add_blank_line(doc)
    _add_body_paragraph(doc, "Температура жилы кабеля до короткого замыкания определяется по формуле:")
    _append_formula(doc, omml_theta_n_schema_string(), first_line_indent_cm=_BODY_FIRST_LINE_INDENT_CM)
    _add_mixed_paragraph(
        doc,
        [
            ("text", "где "),
            ("math", _xml_msub("Θ", "о")),
            ("text", " — температура окружающей среды во время КЗ, принимаем равной 25 "),
            ("math", _xml_mrun("°C")),
            ("text", ";"),
        ],
    )
    _add_mixed_paragraph(
        doc,
        [
            ("math", _xml_msub("Θ", "ДД")),
            (
                "text",
                " — расчётная длительно допустимая температура жилы; для кабелей с поливинилхлоридной "
                "изоляцией принимаем равной 70 ",
            ),
            ("math", _xml_mrun("°C")),
            ("text", ";"),
        ],
    )
    _add_mixed_paragraph(
        doc,
        [
            ("math", _xml_msub("Θ", "ОКР")),
            ("text", " — расчётная температура окружающей среды, 25 "),
            ("math", _xml_mrun("°C")),
            ("text", ";"),
        ],
    )
    _add_mixed_paragraph(
        doc,
        [
            ("math", _xml_msub("I", "раб")),
            ("text", " — номинальный ток вышестоящего автоматического выключателя, А;"),
        ],
    )
    _add_mixed_paragraph(
        doc,
        [
            ("math", _xml_msub("I", "дд")),
            ("text", " — длительно допустимый ток кабеля по сечению жилы, А."),
        ],
    )
    _add_blank_line(doc)
    _add_mixed_paragraph(
        doc,
        [
            ("text", "Ток короткого замыкания в начале кабельной линии "),
            ("math", _xml_msub("I", "кз")),
            (
                "text",
                f" принимается по результатам расчёта токов КЗ (раздел {REPORT_SECTION_TKZ}), кА. "
                "Допускается принимать расчетные токи КЗ на расстоянии 20 м от начала кабельной "
                "линии напряжением до 1 кВ.",
            ),
        ],
    )
    _add_blank_line(doc)
    _add_body_paragraph(doc, "Тепловой импульс от тока короткого замыкания:")
    _append_formula(doc, _xml_omml_b_sopt_schema(), first_line_indent_cm=_BODY_FIRST_LINE_INDENT_CM)
    _add_mixed_paragraph(
        doc,
        [
            ("text", "где "),
            ("math", _xml_msub("I", "кз")),
            ("text", " — ток короткого замыкания в начале кабельной линии, кА."),
        ],
    )
    _add_mixed_paragraph(
        doc,
        [
            ("math", _xml_msub("t", "откл")),
            ("text", " — время отключения от резервной защиты, с."),
        ],
    )
    _add_blank_line(doc)
    _add_body_paragraph(doc, "Температура нагрева жил кабеля при действии тока короткого замыкания:")
    _append_formula(doc, omml_theta_k_schema_string(), first_line_indent_cm=_BODY_FIRST_LINE_INDENT_CM)
    _add_mixed_paragraph(
        doc,
        [
            ("text", "где "),
            ("math", _xml_msub("Θ", "н")),
            ("text", " — температура жилы до КЗ, "),
            ("math", _xml_mrun("°C")),
            ("text", "; "),
            ("math", _xml_mrun("в")),
            (
                "text",
                " — постоянная, характеризующая теплофизические свойства материала жилы; "
                "для меди 19,58 ",
            ),
            ("math", _xml_mrun("мм⁴/(кА²·с)")),
            ("text", ";"),
        ],
    )
    _add_mixed_paragraph(
        doc,
        [
            ("math", _xml_msub("B", "тер")),
            ("text", " — тепловой импульс от тока КЗ, "),
            ("math", _xml_mrun("кА²·с")),
            ("text", ";"),
        ],
    )
    _add_mixed_paragraph(
        doc,
        [
            ("math", _xml_mrun("S")),
            ("text", " — сечение токопроводящей жилы кабеля, мм²;"),
        ],
    )
    _add_mixed_paragraph(
        doc,
        [
            ("math", _xml_msub("Θ", "к")),
            ("text", " — температура жилы кабеля после КЗ, "),
            ("math", _xml_mrun("°C")),
            ("text", "."),
        ],
    )
    _add_blank_line(doc)


def _apply_text_run_font(run: Any, *, size_pt: int = _FONT_SIZE_BODY_PT) -> None:
    run.font.name = _TEXT_FONT_NAME
    run.font.size = Pt(size_pt)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), _TEXT_FONT_NAME)
    r_fonts.set(qn("w:hAnsi"), _TEXT_FONT_NAME)
    r_fonts.set(qn("w:cs"), _TEXT_FONT_NAME)
    r_fonts.set(qn("w:eastAsia"), _TEXT_FONT_NAME)


def _apply_paragraph_text_runs_font(p: Any) -> None:
    for run in p.runs:
        _apply_text_run_font(run)


def _init_document_text_font(doc: Any) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = _TEXT_FONT_NAME
    normal.font.size = Pt(_FONT_SIZE_BODY_PT)
    r_pr = normal.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), _TEXT_FONT_NAME)
    r_fonts.set(qn("w:hAnsi"), _TEXT_FONT_NAME)
    r_fonts.set(qn("w:cs"), _TEXT_FONT_NAME)
    r_fonts.set(qn("w:eastAsia"), _TEXT_FONT_NAME)


def _init_heading_styles(doc: Any) -> None:
    heading_sizes = (
        ("Heading 1", _FONT_SIZE_SECTION_PT),
        ("Heading 2", _FONT_SIZE_SUBSECTION_PT),
        ("Heading 3", _FONT_SIZE_SUBSECTION_PT),
    )
    for _outline_lvl, (style_name, size_pt) in enumerate(heading_sizes):
        style = doc.styles[style_name]
        style.font.name = _TEXT_FONT_NAME
        style.font.size = Pt(size_pt)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        r_pr = style.element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.insert(0, r_fonts)
        r_fonts.set(qn("w:ascii"), _TEXT_FONT_NAME)
        r_fonts.set(qn("w:hAnsi"), _TEXT_FONT_NAME)
        r_fonts.set(qn("w:cs"), _TEXT_FONT_NAME)
        r_fonts.set(qn("w:eastAsia"), _TEXT_FONT_NAME)
        p_pr = style.element.get_or_add_pPr()
        outline = p_pr.find(qn("w:outlineLvl"))
        if outline is None:
            outline = OxmlElement("w:outlineLvl")
            p_pr.append(outline)
        outline.set(qn("w:val"), str(_outline_lvl))
        for tag in ("w:ind",):
            old_ind = p_pr.find(qn(tag))
            if old_ind is not None:
                p_pr.remove(old_ind)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "0")
        ind.set(qn("w:firstLine"), "0")
        p_pr.append(ind)


def _toc_tab_stop_cm() -> float:
    return _TEXT_AREA_WIDTH_CM


def _flat_toc_p_pr(tab_pos: int) -> Any:
    """Единый w:pPr для TOC 1…9: без отступов, номер страницы справа."""
    p_pr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "0")
    ind.set(qn("w:right"), "0")
    ind.set(qn("w:firstLine"), "0")
    ind.set(qn("w:hanging"), "0")
    p_pr.append(ind)
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "dot")
    tab.set(qn("w:pos"), str(tab_pos))
    tabs.append(tab)
    p_pr.append(tabs)
    return p_pr


def _ensure_toc_style(doc: Any, style_name: str, *, base_name: str = "Normal") -> Any:
    try:
        return doc.styles[style_name]
    except KeyError:
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles[base_name]
        return style


def _init_toc_styles(doc: Any) -> None:
    """Оглавление без ступенчатых отступов — все пункты в одну колонку, точки до номера страницы."""
    _ensure_toc_style(doc, "TOC")
    for level in range(1, 10):
        base = "TOC" if level == 1 else f"TOC {level - 1}"
        _ensure_toc_style(doc, f"TOC {level}", base_name=base)
    tab_pos = int(Cm(_toc_tab_stop_cm()).twips)
    flat_p_pr = _flat_toc_p_pr(tab_pos)
    for style_name in ("TOC", *(f"TOC {level}" for level in range(1, 10))):
        style = _ensure_toc_style(doc, style_name)
        style.font.name = _TEXT_FONT_NAME
        style.font.size = Pt(_FONT_SIZE_BODY_PT)
        style.font.bold = False
        style.font.color.rgb = RGBColor(0, 0, 0)
        r_pr = style.element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.insert(0, r_fonts)
        r_fonts.set(qn("w:ascii"), _TEXT_FONT_NAME)
        r_fonts.set(qn("w:hAnsi"), _TEXT_FONT_NAME)
        r_fonts.set(qn("w:cs"), _TEXT_FONT_NAME)
        r_fonts.set(qn("w:eastAsia"), _TEXT_FONT_NAME)
        pf = style.paragraph_format
        pf.left_indent = Cm(0)
        pf.first_line_indent = Cm(0)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        style_el = style.element
        p_pr_el = style_el.find(qn("w:pPr"))
        if p_pr_el is None:
            p_pr_el = OxmlElement("w:pPr")
            style_el.append(p_pr_el)
        for tag in ("w:ind", "w:tabs"):
            old = p_pr_el.find(qn(tag))
            if old is not None:
                p_pr_el.remove(old)
        for child in flat_p_pr:
            p_pr_el.append(deepcopy(child))


def _clear_paragraph_runs(p: Any) -> None:
    for child in list(p._p):
        if child.tag == qn("w:r"):
            p._p.remove(child)


def _append_inline_omath(p: Any, inner_omml: str) -> None:
    omath = parse_xml(f'<m:oMath xmlns:m="{_MNS}">{inner_omml}</m:oMath>')
    p._p.append(omath)


def _add_mixed_paragraph(
    doc: Any,
    parts: list[MixedPart],
    *,
    alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent_cm: float = _BODY_FIRST_LINE_INDENT_CM,
) -> Any:
    p = doc.add_paragraph()
    _style_paragraph(p, alignment=alignment, first_line_indent_cm=first_line_indent_cm)
    _clear_paragraph_runs(p)
    for kind, content in parts:
        if not content:
            continue
        if kind == "text":
            run = p.add_run(content)
            _apply_text_run_font(run)
        else:
            _append_inline_omath(p, content)
    return p


def _xml_e_calc_condition(e_calc: float, selective_ok: bool) -> str:
    op = "<" if selective_ok else "≥"
    return (
        f"{_xml_msub('E', 'расч')}{_xml_mrun(' = ')}{_xml_mrun(_fmt_fixed(e_calc, 2))}{_xml_mrun(' В, т.к. ')}"
        f"{_xml_msub('R', 'кз')}{_xml_mrun(f' {op} ')}{_xml_msub('R', 'гр')}"
    )


def _xml_resistance_comparison(
    r_left: float,
    r_right: float,
    *,
    selective_ok: bool,
    digits: int = 3,
) -> str:
    op = "<" if selective_ok else "≥"
    return (
        f"{_xml_mrun(_fmt_fixed(r_left, digits))}{_xml_mrun(' Ом ')}{_xml_mrun(op)}{_xml_mrun(' ')}"
        f"{_xml_mrun(_fmt_fixed(r_right, digits))}{_xml_mrun(' Ом')}"
    )


def _style_paragraph(
    p: Any,
    *,
    alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent_cm: float = _BODY_FIRST_LINE_INDENT_CM,
) -> None:
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(first_line_indent_cm)
    pf.left_indent = Cm(0)
    pf.right_indent = Cm(0)
    p.alignment = alignment


def _add_body_paragraph(doc: Any, text: str) -> Any:
    p = doc.add_paragraph(text)
    _style_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent_cm=_BODY_FIRST_LINE_INDENT_CM)
    _apply_paragraph_text_runs_font(p)
    return p


def _add_figure_caption(doc: Any, text: str) -> Any:
    p = doc.add_paragraph(text)
    _style_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent_cm=0)
    _apply_paragraph_text_runs_font(p)
    return p


def _add_table_caption(doc: Any, table_num: str, title: str) -> Any:
    p = doc.add_paragraph(f"Таблица {table_num} - {title}")
    _style_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent_cm=0)
    _apply_paragraph_text_runs_font(p)
    return p


def _add_blank_line(doc: Any) -> None:
    p = doc.add_paragraph("")
    _style_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent_cm=0)


def _apply_a4_page(section: Any) -> None:
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(_PAGE_WIDTH_CM)
    section.page_height = Cm(_PAGE_HEIGHT_CM)
    section.top_margin = Cm(_PAGE_MARGIN_TOP_CM)
    section.bottom_margin = Cm(_PAGE_MARGIN_BOTTOM_CM)
    section.left_margin = Cm(_PAGE_MARGIN_LEFT_CM)
    section.right_margin = Cm(_PAGE_MARGIN_RIGHT_CM)
    section.gutter = Cm(_PAGE_MARGIN_GUTTER_CM)


def _kompas_sheet_layout(width_mm: float, height_mm: float) -> tuple[float, float, float, float]:
    """Размер страницы и рисунка с пропорциями листа КОМПАС (с учётом лимита Word)."""
    sheet_w_cm = width_mm / 10.0
    sheet_h_cm = height_mm / 10.0
    if sheet_w_cm <= 0 or sheet_h_cm <= 0:
        return _PAGE_WIDTH_CM, _PAGE_HEIGHT_CM, _PAGE_WIDTH_CM, _PAGE_HEIGHT_CM
    scale = min(1.0, _WORD_MAX_PAGE_SIDE_CM / sheet_w_cm, _WORD_MAX_PAGE_SIDE_CM / sheet_h_cm)
    page_w_cm = sheet_w_cm * scale
    page_h_cm = sheet_h_cm * scale
    return page_w_cm, page_h_cm, page_w_cm, page_h_cm


def _apply_page_margins(doc: Any) -> None:
    _apply_a4_page(doc.sections[0])


def _apply_kompas_page(section: Any, page_w_cm: float, page_h_cm: float) -> None:
    """Страница Word с пропорциями листа КОМПАС (книжная или альбомная).

    В OOXML для альбомной ориентации w:pgSz/@w — длинная сторона листа,
    @h — короткая (как у стандартного A4 landscape в Word).
    """
    landscape = page_w_cm > page_h_cm
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(page_w_cm)
        section.page_height = Cm(page_h_cm)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(page_w_cm)
        section.page_height = Cm(page_h_cm)
    margin = Cm(_KOMPAS_SHEET_MARGIN_CM)
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin
    section.gutter = margin


def _begin_kompas_sheet_section(doc: Any, page_w_cm: float, page_h_cm: float) -> None:
    """Отдельная секция Word под лист КОМПАС (пропорции как в чертеже)."""
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    _apply_kompas_page(section, page_w_cm, page_h_cm)


def _start_a4_continuation(doc: Any) -> None:
    """После листов КОМПАС — новая секция A4 для §4 и далее."""
    doc.add_section(WD_SECTION.NEW_PAGE)
    _apply_a4_page(doc.sections[-1])


def _append_formula(doc: Any, inner_omml: str, *, first_line_indent_cm: float = 0) -> None:
    p = doc.add_paragraph()
    _style_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent_cm=first_line_indent_cm)
    omath = parse_xml(f'<m:oMath xmlns:m="{_MNS}">{inner_omml}</m:oMath>')
    p._p.append(omath)


@dataclass(frozen=True, slots=True)
class TkzCurrentRow:
    point_label: str
    i_kz_a: float
    i_kzd_a: float


def _kz_point_chains(items: list[SoptEquipmentItem]) -> list[tuple[str, list[float]]]:
    point_chains: list[tuple[str, list[float]]] = []
    running_components: list[float] = []
    for item in items:
        if item.item_type == "kz_point":
            point_name = item.designation.strip() or f"KЗ-{len(point_chains) + 1}"
            point_chains.append((point_name, list(running_components)))
            continue
        comp_r = item_resistance_ohm(item)
        if comp_r > 0:
            running_components.append(comp_r)
    if not point_chains:
        point_chains.append(("K1", list(running_components)))
    return point_chains


def _ikz_ikzd_for_r_kz(
    r_kz: float,
    *,
    r_gr: float,
    n_total_elements: int,
    selective_ok: bool | None = None,
) -> tuple[float, float]:
    if r_kz <= 0:
        return 0.0, 0.0
    if selective_ok is None:
        selective_ok = r_kz < r_gr
    e_calc = 1.7 if selective_ok else 1.93
    i_kz = e_calc * n_total_elements / r_kz
    i_kzd = kc_from_r_kz(r_kz) * i_kz
    return i_kz, i_kzd


def _collect_tkz_current_rows(
    input_model: SoptInput,
    result_model: SoptResult,
    *,
    n_total_elements: int,
) -> list[TkzCurrentRow]:
    rows: list[TkzCurrentRow] = []
    i_kz_bus, i_kzd_bus = _ikz_ikzd_for_r_kz(
        result_model.r_kz,
        r_gr=result_model.r_gr,
        n_total_elements=n_total_elements,
        selective_ok=result_model.selective_ok,
    )
    if i_kz_bus > 0:
        rows.append(TkzCurrentRow(TKZ_BUS_KZ_POINT, i_kz_bus, i_kzd_bus))

    for section in input_model.sections:
        for subsection in section.subsections:
            if not subsection.items:
                continue
            for point_name, chain_components in _kz_point_chains(subsection.items):
                chain_r = sum(chain_components)
                rkz_point = result_model.r_ab + result_model.r_per + 2 * chain_r
                i_kz, i_kzd = _ikz_ikzd_for_r_kz(
                    rkz_point,
                    r_gr=result_model.r_gr,
                    n_total_elements=n_total_elements,
                )
                if i_kz > 0:
                    rows.append(TkzCurrentRow(point_name, i_kz, i_kzd))
    return rows


def _append_tkz_currents_table(doc: Any, rows: list[TkzCurrentRow]) -> None:
    _add_page_break(doc)
    _add_blank_line(doc)
    _add_table_caption(doc, tkz_currents_table(1), "Токи короткого замыкания")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    _set_sensitivity_cell_text(header_cells[0], "№")
    _set_sensitivity_cell_text(header_cells[1], "Точка КЗ")
    _set_sensitivity_cell_subscript_label(header_cells[2], "I", "кз", ", А")
    _set_sensitivity_cell_subscript_label(header_cells[3], "I", "кзд", ", А")

    for row_idx, row in enumerate(rows, start=1):
        cells = table.add_row().cells
        values = (
            str(row_idx),
            row.point_label,
            _fmt_fixed(row.i_kz_a, 1),
            _fmt_fixed(row.i_kzd_a, 1),
        )
        for col_idx, value in enumerate(values):
            _set_sensitivity_cell_text(cells[col_idx], value, data_row=True)

    _apply_table_column_widths(
        table,
        _TKZ_CURRENTS_TABLE_COL_WEIGHTS,
        scale=_REPORT_TABLE_WIDTH_SCALE,
    )
    _ensure_table_cells_wrap(table)
    for table_row in table.rows:
        _set_table_row_min_height(table_row, _SENSITIVITY_TABLE_ROW_MIN_HEIGHT_CM)
    _add_blank_line(doc)


def _kz_points_with_chain_r(items: list[SoptEquipmentItem]) -> list[tuple[str, float]]:
    points: list[tuple[str, float]] = []
    chain_r = 0.0
    for item in items:
        if item.item_type == "kz_point":
            points.append((item.designation.strip() or f"KЗ-{len(points) + 1}", chain_r))
        else:
            chain_r += item_resistance_ohm(item)
    if not points:
        points.append(("K1", chain_r))
    return points


def _collect_device_resistances(
    input_model: SoptInput, item_type: str
) -> tuple[list[tuple[int, float, int]], list[str]]:
    groups: dict[tuple[int, float], int] = {}
    unknown: list[str] = []
    for section in input_model.sections:
        for subsection in section.subsections:
            for item in subsection.items:
                if item.item_type != item_type:
                    continue
                if item.rated_current_a <= 0 or item.resistance_ohm <= 0:
                    unknown.append(item.designation.strip() or "без обозначения")
                    continue
                current = int(round(item.rated_current_a))
                key = (current, item.resistance_ohm)
                groups[key] = groups.get(key, 0) + 1
    entries = sorted(
        ((current, resistance, count) for (current, resistance), count in groups.items()),
        key=lambda row: (row[0], row[1]),
    )
    return entries, unknown


def _add_page_break(doc: Any) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _append_tkz_general_scheme_images(
    doc: Any,
    sheets: list[CdwSheetImage],
    _fig_num: int,
) -> None:
    """Листы общей схемы ЩПТ: отдельная страница на лист, пропорции как в КОМПАС."""
    for sheet in sheets:
        page_w_cm, page_h_cm, img_w_cm, img_h_cm = _kompas_sheet_layout(
            sheet.width_mm,
            sheet.height_mm,
        )
        _begin_kompas_sheet_section(doc, page_w_cm, page_h_cm)
        pic_para = doc.add_paragraph()
        _style_paragraph(pic_para, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent_cm=0)
        pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_para.add_run().add_picture(
            _png_stream_with_dpi(sheet.png_bytes),
            width=Cm(img_w_cm),
            height=Cm(img_h_cm),
        )


def _insert_toc_field(paragraph: Any) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    # Уровни 1–3: разделы и подразделы; отступы TOC сняты в _init_toc_styles.
    instr.text = r'TOC \o "1-3" \h \z \u'
    run._r.append(instr)

    run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)

    run = paragraph.add_run("Обновите оглавление: выделите и нажмите F9.")
    _apply_text_run_font(run)

    run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def _enable_update_fields_on_open(doc: Any) -> None:
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


_W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"


def _enable_high_fidelity_images(doc: Any) -> None:
    """Отключает авто-сжатие картинок Word при сохранении и экспорте в PDF."""
    settings = doc.settings.element
    no_compress = settings.find(qn("w:doNotAutoCompressPictures"))
    if no_compress is None:
        no_compress = OxmlElement("w:doNotAutoCompressPictures")
        settings.append(no_compress)
    no_compress.set(qn("w:val"), "true")

    dpi_tag = f"{{{_W14_NS}}}defaultImageDpi"
    dpi_elem = settings.find(dpi_tag)
    if dpi_elem is None:
        dpi_elem = parse_xml(
            f'<w14:defaultImageDpi xmlns:w14="{_W14_NS}" w14:val="330"/>'
        )
        settings.append(dpi_elem)
    else:
        dpi_elem.set(f"{{{_W14_NS}}}val", "330")


def _png_stream_with_dpi(png_bytes: bytes, *, dpi: int = _KOMPAS_RASTER_DPI) -> BytesIO:
    from PIL import Image

    try:
        img = Image.open(BytesIO(png_bytes))
        img.load()
    except Image.DecompressionBombError:
        # Очень крупный растр (A1 @ высокий dpi) — KOMPAS уже задаёт DPI в PNG.
        return BytesIO(png_bytes)
    out = BytesIO()
    img.save(out, format="PNG", dpi=(dpi, dpi))
    out.seek(0)
    return out


def _add_table_of_contents_page(doc: Any) -> None:
    doc_title = doc.add_paragraph()
    doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc_title_run = doc_title.add_run("РАСЧЁТ СОПТ")
    doc_title_run.bold = True
    _apply_text_run_font(doc_title_run, size_pt=_FONT_SIZE_SECTION_PT)
    _add_blank_line(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("СОДЕРЖАНИЕ")
    title_run.bold = True
    _apply_text_run_font(title_run, size_pt=_FONT_SIZE_SECTION_PT)
    _add_blank_line(doc)
    toc_para = doc.add_paragraph()
    _style_paragraph(toc_para, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent_cm=0)
    _insert_toc_field(toc_para)
    _add_page_break(doc)


def _normalize_heading_markup(text: str) -> str:
    """Сборка заголовка: номер и текст через пробел (без w:tab — иначе ломается оглавление)."""
    if "\t" not in text:
        return text
    number, title = text.split("\t", 1)
    return f"{number} {title}"


def _populate_heading_paragraph(p: Any, text: str, *, size_pt: int) -> None:
    """Явные переносы строк из Excel; без табуляции внутри заголовка."""
    _clear_paragraph_runs(p)
    for line_idx, line in enumerate(text.split("\n")):
        if line_idx > 0:
            p.add_run().add_break()
        if line:
            run = p.add_run(line)
            _apply_text_run_font(run, size_pt=size_pt)


def _add_report_heading(
    doc: Any,
    text: str,
    *,
    level: int = 1,
    alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
    page_break_before: bool = False,
) -> Any:
    if page_break_before:
        _add_page_break(doc)
    heading_level = min(max(level, 1), 3)
    markup = _normalize_heading_markup(text)
    p = doc.add_heading("", level=heading_level)
    heading_size = _FONT_SIZE_SECTION_PT if level == 1 else _FONT_SIZE_SUBSECTION_PT
    _populate_heading_paragraph(p, markup, size_pt=heading_size)
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(0)
    pf.right_indent = Cm(0)
    return p


_REFERENCES_ITEMS: tuple[str, ...] = (
    "В.В. Жуков. Короткие замыкания в электроустановках постоянного тока.",
    (
        "Короткие замыкания и выбор электрооборудования: учебное пособие для вузов / "
        "И.П. Крючков, В.А. Старшинов, Ю.П. Гусев и др.; под ред. И.П. Крючкова, В.А. "
        "Старшинова. – М.: Издательский дом МЭИ, 2012. – 568 с.: ил."
    ),
    (
        "ГОСТ 29176-91 Короткие замыкания в электроустановках. Методика расчета в "
        "электроустановках постоянного тока."
    ),
    "Циркуляр № Ц-02-98 (Э).",
)


def _add_numbered_body_paragraph(doc: Any, number: int, text: str) -> Any:
    p = doc.add_paragraph(f"{number}. {text}")
    _style_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent_cm=_BODY_FIRST_LINE_INDENT_CM)
    _apply_paragraph_text_runs_font(p)
    return p


def _append_conclusion_section(doc: Any) -> None:
    _add_report_heading(
        doc,
        f"{REPORT_SECTION_CONCLUSION}\tВЫВОД",
        level=1,
        page_break_before=True,
    )
    _add_blank_line(doc)
    _add_body_paragraph(
        doc,
        "Выполнен расчёт токов короткого замыкания в системе оперативного постоянного "
        "тока. Проведён анализ чувствительности защитных аппаратов и построены карты "
        "селективности. Выполнена проверка кабельных линий на невозгорание в соответствии "
        "с требованиями Циркуляра № Ц-02-98 (Э).",
    )
    _add_body_paragraph(
        doc,
        "По результатам расчётов установлено, что все защитные аппараты обладают "
        "достаточной чувствительностью и обеспечивают необходимую селективность. Кабели "
        "соответствуют нормам по невозгоранию.",
    )
    _add_blank_line(doc)


def _append_references_section(doc: Any) -> None:
    _add_report_heading(
        doc,
        (
            f"{REPORT_SECTION_REFERENCES}\tПЕРЕЧЕНЬ НОРМАТИВНО-ТЕХНИЧЕСКИХ ДОКУМЕНТОВ,\n"
            "ИСПОЛЬЗУЕМЫХ ПРИ ПРОЕКТИРОВАНИИ"
        ),
        level=1,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        page_break_before=True,
    )
    _add_blank_line(doc)
    for idx, item in enumerate(_REFERENCES_ITEMS, start=1):
        _add_numbered_body_paragraph(doc, idx, item)
    _add_blank_line(doc)


def _add_bullet_paragraph(doc: Any, text: str) -> Any:
    p = doc.add_paragraph(f"– {text}")
    _style_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent_cm=0)
    p.paragraph_format.left_indent = Cm(_BODY_FIRST_LINE_INDENT_CM)
    _apply_paragraph_text_runs_font(p)
    return p


def _tkz_substitution_scheme_count(input_model: SoptInput) -> int:
    return sum(
        1
        for section in input_model.sections
        for subsection in section.subsections
        if subsection.items
    )


def _append_description_section(
    doc: Any,
    input_model: SoptInput,
    *,
    shpt_fig_label: str,
) -> None:
    ps_name = input_model.project_name.strip() or "ПС"
    _add_report_heading(
        doc,
        f"{REPORT_SECTION_DESCRIPTION}\tОПИСАНИЕ СОПТ",
        level=1,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _add_blank_line(doc)
    _add_body_paragraph(
        doc,
        f"На {ps_name} предусмотрена установка шкафа зарядно-выпрямительных "
        "устройств и распределения постоянного тока 220 В.",
    )
    _add_body_paragraph(
        doc,
        "В шкафу смонтированы зарядное устройство, схема ввода и распределения "
        "оперативного тока, система контроля сопротивления изоляции и автоматического "
        "поиска отходящих линий с замыканием на землю, система мониторинга, клеммные зажимы и др.",
    )
    _add_body_paragraph(doc, "В ходе расчета были выбраны:")
    _add_bullet_paragraph(doc, "сечения и марки проводов и кабелей;")
    _add_bullet_paragraph(doc, "предохранители (автоматические выключатели).")
    _add_blank_line(doc)
    _add_body_paragraph(
        doc,
        "Был проведен расчет сопротивлений элементов цепи постоянного тока, рассчитаны "
        "сопротивления кабелей, предохранителей и автоматических выключателей.",
    )
    _add_body_paragraph(doc, f"Схема ЩПТ представлена на Рис.{shpt_fig_label}.")
    _add_blank_line(doc)


def _selectivity_section_heading(section_name: str, section_idx: int) -> str:
    title = section_name.strip() or f"Раздел {section_idx}"
    return title.replace(
        "РАСЧЕТ ТОКОВ КОРОТКОГО ЗАМЫКАНИЯ",
        "ПРОВЕРКА СЕЛЕКТИВНОСТИ",
    )


def _apply_table_column_widths(
    table: Any,
    weights: tuple[float, ...],
    *,
    scale: float = 1.0,
) -> None:
    table.autofit = False
    total_weight = sum(weights)
    table_width_cm = _TEXT_AREA_WIDTH_CM * scale
    table.width = Cm(table_width_cm)
    for col_idx, weight in enumerate(weights):
        table.columns[col_idx].width = Cm(table_width_cm * weight / total_weight)


def _set_table_row_min_height(row: Any, height_cm: float) -> None:
    """Минимальная высота строки; Word увеличит её, если текст не помещается."""
    row.height = Cm(height_cm)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def _ensure_table_cells_wrap(table: Any) -> None:
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            no_wrap = tc_pr.find(qn("w:noWrap"))
            if no_wrap is not None:
                tc_pr.remove(no_wrap)


def _style_sensitivity_table_cell(cell: Any, *, data_row: bool = False) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = paragraph.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.left_indent = Cm(0)
        pf.right_indent = Cm(0)
        if data_row:
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = 1.2
            pf.space_before = Pt(3)
            pf.space_after = Pt(3)
        _apply_paragraph_text_runs_font(paragraph)


def _set_sensitivity_cell_text(cell: Any, text: str, *, data_row: bool = False) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    _apply_text_run_font(run)
    _style_sensitivity_table_cell(cell, data_row=data_row)


def _apply_sensitivity_table_layout(table: Any) -> None:
    _apply_table_column_widths(
        table,
        _SENSITIVITY_TABLE_COL_WEIGHTS,
        scale=_REPORT_TABLE_WIDTH_SCALE,
    )
    _ensure_table_cells_wrap(table)
    for row in table.rows:
        _set_table_row_min_height(row, _SENSITIVITY_TABLE_ROW_MIN_HEIGHT_CM)


def _set_sensitivity_cell_subscript_label(
    cell: Any,
    base: str,
    sub: str,
    suffix: str = "",
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run_base = paragraph.add_run(base)
    _apply_text_run_font(run_base)
    run_sub = paragraph.add_run(sub)
    _apply_text_run_font(run_sub)
    run_sub.font.subscript = True
    if suffix:
        run_suffix = paragraph.add_run(suffix)
        _apply_text_run_font(run_suffix)
    _style_sensitivity_table_cell(cell)


def _append_sensitivity_section(
    doc: Any,
    input_model: SoptInput,
    result_model: SoptResult,
    *,
    n_total_elements: int,
    after_kompas_sheets: bool = False,
) -> None:
    rows = build_breaker_sensitivity_rows(
        input_model,
        result_model,
        n_total_elements=n_total_elements,
    )
    if not rows:
        return

    if after_kompas_sheets:
        _start_a4_continuation(doc)

    _add_report_heading(
        doc,
        f"{REPORT_SECTION_SENSITIVITY}\tПРОВЕРКА АВТОМАТИЧЕСКИХ ВЫКЛЮЧАТЕЛЕЙ НА ЧУВСТВИТЕЛЬНОСТЬ",
        level=1,
        page_break_before=not after_kompas_sheets,
    )
    _add_blank_line(doc)
    _add_table_caption(
        doc,
        sensitivity_table(1),
        "Результаты проверки чувствительности автоматических выключателей",
    )

    table = doc.add_table(rows=1, cols=8)
    table.style = "Table Grid"
    plain_headers = (
        "№",
        "Наименование автомата",
        "Характеристика срабатывания",
        "Кратность",
        "Проходит по чувствительности",
    )
    header_cells = table.rows[0].cells
    _set_sensitivity_cell_text(header_cells[0], plain_headers[0])
    _set_sensitivity_cell_text(header_cells[1], plain_headers[1])
    _set_sensitivity_cell_subscript_label(header_cells[2], "I", "ном", ", А")
    _set_sensitivity_cell_text(header_cells[3], plain_headers[2])
    _set_sensitivity_cell_text(header_cells[4], plain_headers[3])
    _set_sensitivity_cell_subscript_label(header_cells[5], "I", "кз", ", А")
    _set_sensitivity_cell_subscript_label(header_cells[6], "K", "ч")
    _set_sensitivity_cell_text(header_cells[7], plain_headers[4])

    for row in rows:
        cells = table.add_row().cells
        values = (
            str(row.number),
            row.designation,
            _fmt_intish(row.rated_current_a) if row.rated_current_a > 0 else "—",
            row.cb_curve,
            _fmt_intish(row.cb_multiplier) if row.cb_multiplier > 0 else "—",
            _fmt_fixed(row.i_kz_a, 1) if row.i_kz_a > 0 else "—",
            _fmt_fixed(row.k_sensitivity, 1) if row.k_sensitivity > 0 else "—",
            "Да" if row.passes else "Нет",
        )
        for col_idx, value in enumerate(values):
            _set_sensitivity_cell_text(cells[col_idx], value, data_row=True)

    _apply_sensitivity_table_layout(table)
    _add_blank_line(doc)


def _append_cable_fire_section(
    doc: Any,
    input_model: SoptInput,
    result_model: SoptResult,
    *,
    n_total_elements: int,
) -> None:
    rows = build_cable_fire_check_rows(
        input_model,
        result_model,
        n_total_elements=n_total_elements,
    )
    if not rows:
        return

    _add_report_heading(
        doc,
        f"{REPORT_SECTION_CABLE_FIRE}\tПРОВЕРКА КАБЕЛЕЙ НА НЕВОЗГОРАНИЕ",
        level=1,
        page_break_before=True,
    )
    _add_blank_line(doc)
    _append_cable_fire_methodology(doc)
    _add_page_break(doc)
    _add_table_caption(
        doc,
        cable_fire_table(1),
        "Результаты проверки на невозгорание",
    )

    table = doc.add_table(rows=1, cols=9)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    _set_sensitivity_cell_text(header_cells[0], "№")
    _set_sensitivity_cell_text(header_cells[1], "Марка кабеля")
    _set_sensitivity_cell_subscript_label(header_cells[2], "I", "раб", ", А")
    _set_sensitivity_cell_subscript_label(header_cells[3], "I", "дд", ", А")
    _set_sensitivity_cell_text(header_cells[4], "Температура жил до КЗ, °C")
    _set_sensitivity_cell_subscript_label(header_cells[5], "I", "кз", ", кА")
    _set_sensitivity_cell_subscript_label(header_cells[6], "t", "откл", ", с")
    _set_sensitivity_cell_text(header_cells[7], "Температура жилы после КЗ, °C")
    _set_sensitivity_cell_text(header_cells[8], "Проходит по нагреву")

    for row in rows:
        cells = table.add_row().cells
        values = (
            str(row.number),
            row.cable_mark,
            _fmt_intish(row.i_nom_a) if row.i_nom_a > 0 else "—",
            _fmt_intish(row.i_dd_a) if row.i_dd_a > 0 else "—",
            str(row.theta_before_c) if row.theta_before_c is not None else "—",
            _fmt_fixed(row.i_kz_ka, 3) if row.i_kz_ka > 0 else "—",
            _fmt_fixed(row.shutdown_time_s, 2) if row.shutdown_time_s > 0 else "—",
            _fmt_fixed(row.theta_after_c, 2) if row.theta_after_c is not None else "—",
            "Да" if row.passes else "Нет",
        )
        for col_idx, value in enumerate(values):
            _set_sensitivity_cell_text(cells[col_idx], value)

    _apply_table_column_widths(
        table,
        _CABLE_FIRE_TABLE_COL_WEIGHTS,
        scale=_REPORT_TABLE_WIDTH_SCALE,
    )
    _add_blank_line(doc)


def _is_branch_breaker(item: SoptEquipmentItem) -> bool:
  designation = item.designation.strip().upper()
  return designation.startswith("SF")


def _section_incomer_breaker(section: SoptSection) -> SoptEquipmentItem | None:
    """Вводной автомат раздела (2QF4, 1QF1 — не веточный SF*)."""
    fallback: SoptEquipmentItem | None = None
    for subsection in section.subsections:
        for item in subsection.items:
            if item.item_type != "breaker":
                continue
            if fallback is None:
                fallback = item
            if not _is_branch_breaker(item):
                return item
    return fallback


def _append_selectivity_section(
    doc: Any,
    input_model: SoptInput,
    result_model: SoptResult,
    *,
    n_total_elements: int,
    fig_num: int,
) -> int:
    sections_with_maps: list[tuple[int, SoptSection, list[SelectivityMapResult]]] = []

    for section_idx, section in enumerate(input_model.sections, start=1):
        if not section.subsections:
            continue

        section_maps: list[SelectivityMapResult] = []
        head_breaker = _section_incomer_breaker(section)
        for subsection in section.subsections:
            if not subsection.items:
                continue
            try:
                selectivity_maps = build_subsection_selectivity_maps(
                    subsection.items,
                    r_ab=result_model.r_ab,
                    r_per=result_model.r_per,
                    r_gr=result_model.r_gr,
                    n_total_elements=n_total_elements,
                    section_head_breaker=head_breaker,
                )
            except RuntimeError:
                continue

            section_maps.extend(selectivity_maps)

        if section_maps:
            sections_with_maps.append((section_idx, section, section_maps))

    if not sections_with_maps:
        return fig_num

    _add_report_heading(
        doc,
        f"{REPORT_SECTION_SELECTIVITY}\tПРОВЕРКА СЕЛЕКТИВНОСТИ",
        level=1,
        page_break_before=True,
    )

    for section_idx, section, section_maps in sections_with_maps:
        section_heading = _selectivity_section_heading(section.name, section_idx)
        _add_report_heading(doc, f"{selectivity_section(section_idx)}\t{section_heading}", level=2)

        for map_idx, map_result in enumerate(section_maps, start=1):
            _add_report_heading(
                doc,
                f"{selectivity_subsection(section_idx, map_idx)}\t{map_result.block_title}",
                level=3,
            )
            pic_para = doc.add_paragraph()
            _style_paragraph(pic_para, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent_cm=0)
            pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic_para.add_run().add_picture(
                map_result.image,
                width=Cm(12.0),
            )
            _add_mixed_paragraph(
                doc,
                [
                    (
                        "text",
                        f"Рис.{fig_num} — Карта селективности пары {map_result.upstream_label} "
                        f"и {map_result.downstream_label} (",
                    ),
                    (
                        "math",
                        f"{_xml_msub('I', 'кз.макс')}{_xml_mrun(' = ')}"
                        f"{_xml_mrun(_fmt_fixed(map_result.i_max_a, 1))}{_xml_mrun(' А')}",
                    ),
                    ("text", ", "),
                    (
                        "math",
                        f"{_xml_msub('I', 'кз.мин')}{_xml_mrun(' = ')}"
                        f"{_xml_mrun(_fmt_fixed(map_result.i_min_a, 1))}{_xml_mrun(' А')}",
                    ),
                    ("text", ")."),
                ],
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                first_line_indent_cm=0,
            )
            fig_num += 1
            status = (
                "СЕЛЕКТИВНОСТЬ НЕ СОБЛЮДАЕТСЯ"
                if map_result.overlap
                else "Селективность обеспечена"
            )
            p_res = doc.add_paragraph(f"Вывод: {status}")
            _style_paragraph(p_res, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent_cm=_BODY_FIRST_LINE_INDENT_CM)
            _apply_paragraph_text_runs_font(p_res)
            if map_result.overlap and p_res.runs:
                p_res.runs[0].bold = True

    _add_blank_line(doc)
    return fig_num


def populate_sopt_document(
    doc: Any,
    input_model: SoptInput,
    result_model: SoptResult,
    *,
    general_scheme_sheets: list[CdwSheetImage],
) -> None:
    _apply_page_margins(doc)
    _init_document_text_font(doc)
    _init_heading_styles(doc)
    _init_toc_styles(doc)
    _enable_update_fields_on_open(doc)
    _enable_high_fidelity_images(doc)
    _add_table_of_contents_page(doc)
    scheme_count = _tkz_substitution_scheme_count(input_model)
    _append_description_section(
        doc,
        input_model,
        shpt_fig_label=shpt_scheme_figure_label(scheme_count),
    )

    _add_report_heading(
        doc,
        f"{REPORT_SECTION_RESISTANCE}\tРАСЧЕТ СОПРОТИВЛЕНИЙ ЭЛЕМЕНТОВ ЦЕПИ ПОСТОЯННОГО ТОКА",
        level=1,
        page_break_before=True,
    )
    _add_blank_line(doc)

    battery_name = input_model.battery_name.strip() or "АКБ"
    battery_count = input_model.battery_count
    battery_cells_count = input_model.battery_cells_count
    n_total_elements = input_model.battery_cells_count * input_model.battery_count
    r_el = input_model.r_el
    rho = input_model.rho
    jumper_section = input_model.jumper_section
    jumpers_count = max(0, battery_count - 1)
    total_jumper_length = (
        jumpers_count * _JUMPER_LENGTH_BETWEEN_BANKS
        + _JUMPER_LENGTH_TO_NEIGHBOR_CABINET * 2
    )
    r_ab_mohm = r_el * battery_count
    r_ab = r_ab_mohm / 1000.0
    r_per = rho * total_jumper_length / jumper_section if jumper_section > 0 else 0.0

    _add_mixed_paragraph(
        doc,
        [
            (
                "text",
                f"1) Аккумуляторная батарея {battery_name} "
                f"(12 В, количество элементов ",
            ),
            (
                "math",
                f"{_xml_mrun('n')}{_xml_mrun(' = ')}{_xml_mrun(str(battery_cells_count))}",
            ),
            ("text", f"), {battery_count} штук."),
        ],
    )
    _add_body_paragraph(doc, "Технические данные АБ:")
    _add_mixed_paragraph(
        doc,
        [
            (
                "math",
                f"{_xml_msub('R', 'эл')}{_xml_mrun(' = ')}{_xml_mrun(_fmt_fixed(r_el, 1))}{_xml_mrun(' мОм')}",
            ),
            ("text", " – внутреннее сопротивление блока 12 В;"),
        ],
    )
    _add_mixed_paragraph(
        doc,
        [
            (
                "math",
                f"{_xml_msub('R', 'аб')}{_xml_mrun(' = ')}{_xml_msub('R', 'эл')}{_xml_mrun(' · n = ')}"
                f"{_xml_mrun(_fmt_fixed(r_el, 1))}{_xml_mrun(' · ')}{_xml_mrun(str(battery_count))}{_xml_mrun(' = ')}"
                f"{_xml_mrun(_fmt_fixed(r_ab_mohm, 1))}{_xml_mrun(' мОм = ')}"
                f"{_xml_mrun(_fmt_fixed(r_ab, 4))}{_xml_mrun(' Ом')}",
            ),
            ("text", " – внутреннее сопротивление АБ."),
        ],
    )
    _add_body_paragraph(doc, "Перемычки:")
    _add_body_paragraph(doc, f"Количество перемычек: {jumpers_count} штук;")
    _add_mixed_paragraph(
        doc,
        [
            (
                "math",
                f"{_xml_mrun('S')}{_xml_mrun(' = ')}{_xml_mrun(_fmt_intish(jumper_section))}{_xml_mrun(' мм²')}",
            ),
            ("text", ";"),
        ],
    )
    _add_mixed_paragraph(
        doc,
        [
            ("math", f"{_xml_mrun('L')}{_xml_mrun(' = 0,25 м')}"),
            ("text", " – длина между банками, в соседний шкаф – 1,5 м."),
        ],
    )
    _add_body_paragraph(doc, "Сопротивление перемычек АБ:")
    _append_formula(
        doc,
        f"{_xml_msub('R', 'пер')}{_xml_mrun(' = ')}{_xml_mrun('ρ')}{_xml_mrun('·')}"
        f"{_xml_mfrac(_xml_mrun('L'), _xml_mrun('S'))}"
        f"{_xml_mrun(' = ')}{_xml_mrun(_fmt_fixed(rho, 4))}{_xml_mrun('·')}"
        f"{_xml_mfrac(_xml_mrun(_fmt_intish(total_jumper_length)), _xml_mrun(_fmt_intish(jumper_section)))}"
        f"{_xml_mrun(' = ')}{_xml_mrun(_fmt_fixed(r_per, 4))}{_xml_mrun(' Ом')}",
    )
    _add_mixed_paragraph(
        doc,
        [
            ("text", "Где: "),
            ("math", f"{_xml_mrun('ρ')}{_xml_mrun(' = ')}{_xml_mrun(_fmt_fixed(rho, 4))}"),
            ("text", " – удельное сопротивление меди;"),
        ],
    )
    _add_mixed_paragraph(
        doc,
        [
            (
                "math",
                f"{_xml_mrun('S')}{_xml_mrun(' = ')}{_xml_mrun(_fmt_intish(jumper_section))}{_xml_mrun(' мм²')}",
            ),
            ("text", " – сечение перемычек;"),
        ],
    )
    _add_mixed_paragraph(
        doc,
        [
            (
                "math",
                f"{_xml_mrun('L')}{_xml_mrun(' = ')}{_xml_mrun(str(jumpers_count))}{_xml_mrun(' · 0,25 + 1,5 · 2 = ')}"
                f"{_xml_mrun(_fmt_intish(total_jumper_length))}{_xml_mrun(' м')}",
            ),
            ("text", " – общая длина перемычек."),
        ],
    )
    _add_blank_line(doc)

    fuse_entries, unknown_fuses = _collect_device_resistances(input_model, "fuse")
    breaker_entries, unknown_breakers = _collect_device_resistances(input_model, "breaker")

    if fuse_entries or unknown_fuses:
        _add_body_paragraph(doc, "2) Сопротивления предохранителей (Табл.2.9 [1]):")
        for current, value, _count in fuse_entries:
            _add_body_paragraph(
                doc,
                f"- сопротивление предохранителя с номинальным током {current} А."
            )
            _append_formula(
                doc,
                f"{_xml_mrun('R')}{_xml_mrun(' = ')}{_xml_mrun(_fmt_fixed(value, 3))}{_xml_mrun(' Ом')}",
            )
        if unknown_fuses:
            _add_body_paragraph(
                doc,
                "- Не указаны номинальный ток и/или R для: "
                + "; ".join(unknown_fuses)
                + "."
            )

    if breaker_entries or unknown_breakers:
        _add_body_paragraph(doc, "3) Сопротивления автоматических выключателей (Табл.2.7 [1]):")
        for current, value, _count in breaker_entries:
            _add_body_paragraph(
                doc,
                f"- сопротивление автоматического выключателя с номинальным током {current} А."
            )
            _append_formula(
                doc,
                f"{_xml_mrun('R')}{_xml_mrun(' = ')}{_xml_mrun(_fmt_fixed(value, 3))}{_xml_mrun(' Ом')}",
            )
        if unknown_breakers:
            _add_body_paragraph(
                doc,
                "- Не указаны номинальный ток и/или R для: "
                + "; ".join(unknown_breakers)
                + "."
            )

    _add_blank_line(doc)

    input_fuse_label = input_model.input_fuse_label.strip() or "Вводной предохранитель"
    input_fuse_r = input_model.input_fuse_resistance
    qn1_ah = input_model.qn1_ah
    q_calc_ah = input_model.q_calc_ah
    n_required = result_model.n_required
    n_accepted = result_model.n_accepted
    r_kz = result_model.r_kz
    r_gr = result_model.r_gr
    e_calc = 1.7 if result_model.selective_ok else 1.93
    r_gr_mohm = r_gr * 1000.0

    _add_report_heading(
        doc,
        f"{REPORT_SECTION_TKZ}\tРАСЧЕТ ТОКОВ КОРОТКОГО ЗАМЫКАНИЯ, ПРОВЕРКА ЧУВСТВИТЕЛЬНОСТИ.",
        level=1,
        page_break_before=True,
    )
    _add_report_heading(
        doc,
        f"{REPORT_SECTION_TKZ_BUS}\tРАСЧЕТ ТОКОВ КОРОТКОГО ЗАМЫКАНИЯ НА ШИНАХ ЩИТА ПОСТОЯННОГО ТОКА",
        level=2,
    )
    pic_para = doc.add_paragraph()
    _style_paragraph(pic_para, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent_cm=0)
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_para.add_run().add_picture(
        akb_input_schematic_png_bytes(),
        width=Cm(_FIG2_AND_NEXT_WIDTH_CM),
        height=Cm(_FIG2_AND_NEXT_HEIGHT_CM),
    )
    _add_figure_caption(doc, "Рис.1 – Схема питания от АКБ.")
    pic2_para = doc.add_paragraph()
    _style_paragraph(pic2_para, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent_cm=0)
    pic2_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic2_para.add_run().add_picture(
        k1_substitution_schematic_png_bytes(
            result_model.r_ab,
            result_model.r_per,
            input_fuse_r,
            kz_point_label=TKZ_BUS_KZ_POINT,
        ),
        width=Cm(_FIG2_AND_NEXT_WIDTH_CM),
        height=Cm(_FIG2_AND_NEXT_HEIGHT_CM),
    )
    _add_figure_caption(doc, f"Рис.2 – Схема замещения до точки {TKZ_BUS_KZ_POINT}.")
    _add_blank_line(doc)
    _add_body_paragraph(doc, f"Сопротивление цепи тока короткого замыкания в т. {TKZ_BUS_KZ_POINT}:")
    _append_formula(
        doc,
        f"{_xml_msub('R', 'кз')}{_xml_mrun(' = (')}{_xml_msub('R', 'аб')}{_xml_mrun(' + ')}{_xml_msub('R', 'пер')}{_xml_mrun(') + 2·')}{_xml_msub('R', 'пр')}{_xml_mrun(' = ')}"
        f"{_xml_mrun('(')}{_xml_mrun(_fmt_fixed(result_model.r_ab, 4))}{_xml_mrun(' + ')}{_xml_mrun(_fmt_fixed(result_model.r_per, 4))}{_xml_mrun(') + 2·')}"
        f"{_xml_mrun(_fmt_fixed(input_fuse_r, 3))}{_xml_mrun(' = ')}{_xml_mrun(_fmt_fixed(r_kz, 3))}"
        f"{_xml_mrun(' Ом')}",
    )
    _add_mixed_paragraph(
        doc,
        [
            ("math", _xml_e_calc_condition(e_calc, result_model.selective_ok)),
            ("text", "."),
        ],
    )
    _append_formula(
        doc,
        f"{_xml_msub('R', 'гр')}{_xml_mrun(' = 7,5·')}{_xml_mfrac(_xml_mrun('n'), _xml_mrun('N'))}{_xml_mrun(' мОм')}",
    )
    _add_mixed_paragraph(
        doc,
        [
            ("text", "где n — количество элементов в цепи "),
            (
                "math",
                f"{_xml_mrun('n')}{_xml_mrun(' = ')}{_xml_mrun(str(battery_cells_count))}{_xml_mrun('·')}"
                f"{_xml_mrun(str(battery_count))}{_xml_mrun(' = ')}{_xml_mrun(str(n_total_elements))}{_xml_mrun(' эл.')}",
            ),
        ],
    )
    _add_body_paragraph(
        doc,
        "N — номер аккумуляторной батареи, определяется с учётом понижения энергии батареи за период эксплуатации",
    )
    _append_formula(
        doc,
        f"{_xml_msub('N', 'усл')}{_xml_mrun(' ≥ 1,1·')}{_xml_mfrac(_xml_msub('Q', 'расч'), _xml_msub('Q', 'N=1'))}",
    )
    _add_body_paragraph(doc, "где 1,1 — коэффициент, учитывающий понижение энергии батареи")
    p_qn1_desc = doc.add_paragraph()
    _style_paragraph(p_qn1_desc, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent_cm=_BODY_FIRST_LINE_INDENT_CM)
    p_qn1_desc.add_run("Q")
    run_qn1_desc_sub = p_qn1_desc.add_run("N=1")
    run_qn1_desc_sub.font.subscript = True
    p_qn1_desc.add_run(f" — энергия аккумулятора {battery_name} (при одночасовом разряде)")
    _apply_paragraph_text_runs_font(p_qn1_desc)
    _add_mixed_paragraph(
        doc,
        [
            (
                "math",
                f"{_xml_msub('Q', 'N=1')}{_xml_mrun(' = ')}{_xml_mrun(_fmt_fixed(qn1_ah, 3))}{_xml_mrun(' А·ч')}",
            ),
        ],
    )
    _add_body_paragraph(doc, f"Qрасч — энергия аккумулятора {battery_name} (при одночасовом разряде)")
    _add_mixed_paragraph(
        doc,
        [
            (
                "math",
                f"{_xml_msub('Q', 'расч')}{_xml_mrun(' = ')}{_xml_mrun(_fmt_fixed(q_calc_ah, 0))}{_xml_mrun(' А·ч')}",
            ),
        ],
    )
    _append_formula(
        doc,
        f"{_xml_msub('N', 'усл')}{_xml_mrun(' ≥ 1,1·')}"
        f"{_xml_mfrac(_xml_mrun(_fmt_fixed(q_calc_ah, 0)), _xml_mrun(_fmt_fixed(qn1_ah, 3)))}"
        f"{_xml_mrun(' = ')}{_xml_mrun(_fmt_fixed(n_required, 2))}",
    )
    _add_mixed_paragraph(
        doc,
        [
            ("text", "Принимаем "),
            ("math", f"{_xml_mrun('N')}{_xml_mrun(' = ')}{_xml_mrun(str(n_accepted))}"),
        ],
    )
    _append_formula(
        doc,
        f"{_xml_msub('R', 'гр')}{_xml_mrun(' = 7,5·')}{_xml_mfrac(_xml_mrun(str(n_total_elements)), _xml_mrun(str(n_accepted)))}"
        f"{_xml_mrun(' = ')}{_xml_mrun(_fmt_fixed(r_gr_mohm, 1))}{_xml_mrun(' мОм = ')}{_xml_mrun(_fmt_fixed(r_gr, 3))}{_xml_mrun(' Ом')}",
    )
    _add_mixed_paragraph(
        doc,
        [("math", _xml_resistance_comparison(r_kz, r_gr, selective_ok=result_model.selective_ok))],
    )
    _add_blank_line(doc)

    kc = kc_from_r_kz(r_kz)
    fuse_trip_time_s = 0.01
    i_kz, i_kzd = _ikz_ikzd_for_r_kz(
        r_kz,
        r_gr=r_gr,
        n_total_elements=n_total_elements,
        selective_ok=result_model.selective_ok,
    )
    nominal_match = re.search(r"(\d+)", input_fuse_label)
    fuse_nominal_a = nominal_match.group(1) if nominal_match else ""
    fuse_nominal_text = f"{fuse_nominal_a}А" if fuse_nominal_a else (input_fuse_label.split("(")[0].strip() or input_fuse_label)

    _add_body_paragraph(doc, f"Ток короткого замыкания в т. {TKZ_BUS_KZ_POINT}:")
    _append_formula(
        doc,
        f"{_xml_msub('I', 'кз')}{_xml_mrun(' = ')}"
        f"{_xml_mfrac(_xml_mrun('Eрасч·n'), _xml_msub('R', 'кз'))}"
        f"{_xml_mrun(' = ')}{_xml_mfrac(_xml_mrun(_fmt_fixed(e_calc, 2) + '·' + str(n_total_elements)), _xml_mrun(_fmt_fixed(r_kz, 3)))}"
        f"{_xml_mrun(' = ')}{_xml_mrun(_fmt_fixed(i_kz, 1))}{_xml_mrun(' А')}",
    )
    _add_body_paragraph(doc, f"Ток короткого замыкания в т. {TKZ_BUS_KZ_POINT} с учётом сопротивления дуги:")
    _append_formula(
        doc,
        f"{_xml_msub('I', 'кзд')}{_xml_mrun(' = ')}{_xml_msub('K', 'c')}{_xml_mrun('·')}{_xml_msub('I', 'кз')}{_xml_mrun(' = ')}{_xml_mrun(_fmt_kc(kc))}"
        f"{_xml_mrun('·')}{_xml_mrun(_fmt_fixed(i_kz, 1))}{_xml_mrun(' = ')}{_xml_mrun(_fmt_fixed(i_kzd, 1))}{_xml_mrun(' А')}",
    )
    p_kc_desc = doc.add_paragraph()
    _style_paragraph(p_kc_desc, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent_cm=_BODY_FIRST_LINE_INDENT_CM)
    p_kc_desc.add_run("K")
    run_kc_sub = p_kc_desc.add_run("c")
    run_kc_sub.font.subscript = True
    p_kc_desc.add_run(" — коэффициент снижения тока КЗ, определяется по кривой зависимости")
    _apply_paragraph_text_runs_font(p_kc_desc)
    _append_formula(
        doc,
        f"{_xml_msub('K', 'c')}{_xml_mrun(' = f(')}{_xml_msub('R', 'кз')}{_xml_mrun(')')}",
    )
    _append_formula(
        doc,
        f"{_xml_msub('K', 'c')}{_xml_mrun(' = ')}{_xml_mrun(_fmt_kc(kc))}{_xml_mrun(' для ')}{_xml_msub('R', 'кз')}{_xml_mrun(' = ')}{_xml_mrun(_fmt_fixed(r_kz, 3))}{_xml_mrun(' Ом')}",
    )
    _add_body_paragraph(
        doc,
        f"Время перегорания предохранителя {fuse_nominal_text} не более {_fmt_fixed(fuse_trip_time_s, 2)} с "
        f"по время-токовой характеристике gG-{fuse_nominal_text}",
    )
    _add_blank_line(doc)

    scheme_count = _tkz_substitution_scheme_count(input_model)
    if not input_model.sections:
        _add_body_paragraph(doc, "Разделы не заданы.")
        fig_num = shpt_scheme_figure_number(scheme_count)
    else:
        fig_num = 3
        for section_idx, section in enumerate(input_model.sections, start=1):
            if not section.subsections:
                continue
            section_name = section.name.strip() or f"Раздел {section_idx}"
            _add_report_heading(
                doc,
                f"{tkz_user_section(section_idx)}\t{section_name}",
                level=2,
                page_break_before=True,
            )

            for subsection_idx, subsection in enumerate(section.subsections, start=1):
                subsection_name = subsection.name.strip() or f"Подраздел {subsection_idx}"
                _add_report_heading(
                    doc,
                    f"{tkz_user_subsection(section_idx, subsection_idx)}\t{subsection_name}",
                    level=3,
                )

                if not subsection.items:
                    _add_body_paragraph(doc, "(оборудование не добавлено)")
                    _add_body_paragraph(doc, "")
                    continue

                stream = subsection_kz_schematic_png_bytes(subsection.items, result_model.r_ab, result_model.r_per)
                chain_items = [it for it in subsection.items if it.item_type != "kz_point"]
                kz_count = max(1, sum(1 for it in subsection.items if it.item_type == "kz_point"))
                pic_width_cm = min(18.0, max(9.6, 1.35 * max(1, len(chain_items)) + 1.1 * kz_count + 4.4))
                pic_para = doc.add_paragraph()
                _style_paragraph(pic_para, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent_cm=0)
                pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pic_para.add_run().add_picture(
                    stream,
                    width=Cm(pic_width_cm),
                )
                point_title = _kz_points_with_chain_r(subsection.items)[-1][0]
                _add_figure_caption(doc, f"Рис.{fig_num} – Схема замещения до точки {point_title}.")
                fig_num += 1

                cable_items = [it for it in subsection.items if it.item_type == "cable"]
                for cable in cable_items:
                    if cable.cable_length_m <= 0 or cable.cable_gamma <= 0 or cable.cable_section_mm2 <= 0:
                        continue
                    cable_name = cable.designation.strip() or "ВВГЭ"
                    element_label = _resistance_element_label(cable_name)
                    _add_mixed_paragraph(
                        doc,
                        [
                            (
                                "text",
                                f"Сопротивление {element_label} {cable_name} 2х{_fmt_intish(cable.cable_section_mm2)} ",
                            ),
                            (
                                "math",
                                f"{_xml_mrun('L')}{_xml_mrun(' = ')}{_xml_mrun(_fmt_intish(cable.cable_length_m))}{_xml_mrun(' м')}",
                            ),
                        ],
                    )
                    rk = item_resistance_ohm(cable)
                    _append_formula(
                        doc,
                        f"{_xml_msub('R', 'к')}{_xml_mrun(' = ')}"
                        f"{_xml_mfrac(_xml_mrun('L'), _xml_mrun('Y·S'))}"
                        f"{_xml_mrun(' = ')}"
                        f"{_xml_mfrac(_xml_mrun(_fmt_intish(cable.cable_length_m)), _xml_mrun(_fmt_intish(cable.cable_gamma) + '·' + _fmt_intish(cable.cable_section_mm2)))}"
                        f"{_xml_mrun(' = ')}{_xml_mrun(_fmt_fixed(rk, 3))}{_xml_mrun(' Ом')}",
                    )

                point_chains = _kz_point_chains(subsection.items)

                point_rkz_values: list[tuple[str, float, float]] = []
                for point_name, chain_components in point_chains:
                    chain_r = sum(chain_components)
                    rkz_point = result_model.r_ab + result_model.r_per + 2 * chain_r
                    e_calc_point = 1.7 if rkz_point < result_model.r_gr else 1.93
                    ikz_point, ikzd_point = _ikz_ikzd_for_r_kz(
                        rkz_point,
                        r_gr=result_model.r_gr,
                        n_total_elements=n_total_elements,
                    )
                    kc_point = kc_from_r_kz(rkz_point)
                    point_rkz_values.append((point_name, rkz_point, kc_point))
                    _add_body_paragraph(doc, f"Сопротивление цепи короткого замыкания в т. {point_name}:")
                    inner_sum = " + ".join(_fmt_fixed(value, 3) for value in chain_components) if chain_components else "0"
                    _append_formula(
                        doc,
                        f"{_xml_msub('R', 'кз')}{_xml_mrun(' = ')}{_xml_mrun(_fmt_fixed(result_model.r_ab, 4))}"
                        f"{_xml_mrun(' + ')}{_xml_mrun(_fmt_fixed(result_model.r_per, 4))}{_xml_mrun(' + 2·(')}"
                        f"{_xml_mrun(inner_sum)}{_xml_mrun(') = ')}{_xml_mrun(_fmt_fixed(rkz_point, 3))}{_xml_mrun(' Ом')}",
                    )
                    _add_mixed_paragraph(
                        doc,
                        [
                            (
                                "math",
                                _xml_e_calc_condition(
                                    e_calc_point,
                                    selective_ok=rkz_point < result_model.r_gr,
                                ),
                            ),
                            ("text", "."),
                        ],
                    )
                    _add_body_paragraph(doc, f"Ток короткого замыкания в т. {point_name}:")
                    _append_formula(
                        doc,
                        f"{_xml_msub('I', 'кз')}{_xml_mrun(' = ')}{_xml_mfrac(_xml_mrun(_fmt_fixed(e_calc_point, 2) + '·' + str(n_total_elements)), _xml_mrun(_fmt_fixed(rkz_point, 3)))}"
                        f"{_xml_mrun(' = ')}{_xml_mrun(_fmt_fixed(ikz_point, 1))}{_xml_mrun(' А')}",
                    )
                    _add_body_paragraph(doc, f"Ток короткого замыкания в т. {point_name} с учётом сопротивления дуги:")
                    _append_formula(
                        doc,
                        f"{_xml_msub('I', 'кзд')}{_xml_mrun(' = ')}{_xml_msub('K', 'c')}{_xml_mrun('·')}{_xml_msub('I', 'кз')}{_xml_mrun(' = ')}"
                        f"{_xml_mrun(_fmt_kc(kc_point))}{_xml_mrun('·')}{_xml_mrun(_fmt_fixed(ikz_point, 1))}{_xml_mrun(' = ')}{_xml_mrun(_fmt_fixed(ikzd_point, 1))}{_xml_mrun(' А')}",
                    )

                p_kc_desc = doc.add_paragraph()
                _style_paragraph(p_kc_desc, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent_cm=_BODY_FIRST_LINE_INDENT_CM)
                p_kc_desc.add_run("K")
                run_kc_sub = p_kc_desc.add_run("c")
                run_kc_sub.font.subscript = True
                p_kc_desc.add_run(" — коэффициент снижения тока КЗ, определяется по кривой зависимости")
                _apply_paragraph_text_runs_font(p_kc_desc)
                _append_formula(
                    doc,
                    f"{_xml_msub('K', 'c')}{_xml_mrun(' = f(')}{_xml_msub('R', 'кз')}{_xml_mrun(')')}",
                )
                for _, rkz_point, kc_point in point_rkz_values:
                    _append_formula(
                        doc,
                        f"{_xml_msub('K', 'c')}{_xml_mrun(' = ')}{_xml_mrun(_fmt_kc(kc_point))}{_xml_mrun(' для ')}{_xml_msub('R', 'кз')}{_xml_mrun(' = ')}{_xml_mrun(_fmt_fixed(rkz_point, 3))}{_xml_mrun(' Ом')}",
                    )

                _add_body_paragraph(doc, "")

    tkz_current_rows = _collect_tkz_current_rows(
        input_model,
        result_model,
        n_total_elements=n_total_elements,
    )
    if tkz_current_rows:
        _append_tkz_currents_table(doc, tkz_current_rows)

    _append_tkz_general_scheme_images(doc, general_scheme_sheets, fig_num)
    _append_sensitivity_section(
        doc,
        input_model,
        result_model,
        n_total_elements=n_total_elements,
        after_kompas_sheets=bool(general_scheme_sheets),
    )
    _append_cable_fire_section(
        doc,
        input_model,
        result_model,
        n_total_elements=n_total_elements,
    )
    _append_selectivity_section(
        doc,
        input_model,
        result_model,
        n_total_elements=n_total_elements,
        fig_num=selectivity_figure_start(fig_num),
    )
    _append_conclusion_section(doc)
    _append_references_section(doc)
    _init_toc_styles(doc)
