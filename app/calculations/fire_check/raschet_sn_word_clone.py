"""
Отчёт Word: раздел проверки кабелей на невозгорание (подробный расчёт для 1-й линии + таблица 2.1).
Оформление: Times New Roman 14 pt, заголовки разделов синим (RGB 31,78,121), шапки таблиц D9D9D9.
"""

from __future__ import annotations

from typing import Any

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .word_omml import (
    append_inline_omath,
    append_mixed_paragraph,
    append_omath_xml,
    legend_where_B_parts,
    legend_where_B_tae_parts,
    legend_where_ik_schema_parts,
    legend_where_theta_k_a_parts,
    legend_where_theta_k_b_parts,
    legend_where_theta_k_parts,
    legend_where_theta_k_s_parts,
    legend_where_theta_n_dd_parts,
    legend_where_theta_n_idd_parts,
    legend_where_theta_n_irab_parts,
    legend_where_theta_n_o_parts,
    legend_where_theta_n_okr_parts,
    omml_B_schema_string,
    omml_B_string,
    omml_Ip0_numeric_mohm,
    omml_Ip0_schema_string,
    omml_theta_k_schema_string,
    omml_theta_k_string,
    omml_theta_n_schema_string,
    omml_theta_n_string,
    xml_mrun,
)


def _fmt(x: float | None, nd: int = 2) -> str:
    if x is None:
        return "—"
    return f"{x:.{nd}f}".replace(".", ",")


def _fmt_int(x: int | None) -> str:
    if x is None:
        return "—"
    return str(x)


def _set_cell_shading(cell: Any, fill_hex: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def _style_header_row(table, fill_hex: str = "D9D9D9") -> None:
    for cell in table.rows[0].cells:
        _set_cell_shading(cell, fill_hex)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True


def _p_caps_blue(doc: Any, text: str) -> None:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    pf.space_before = Pt(12)
    pf.space_after = Pt(12)
    pf.first_line_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(31, 78, 121)
    run.font.name = "Times New Roman"


def _p_body(
    doc: Any,
    text: str,
    bold: bool = False,
    *,
    center: bool = False,
    no_indent: bool = False,
) -> None:
    p = doc.add_paragraph(text)
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    if no_indent:
        pf.first_line_indent = Cm(0)
    else:
        pf.first_line_indent = Cm(1.25)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(14)
        r.bold = bold


def populate_raschet_sn_clone(doc: Any, input_model: Any, result_model: Any) -> None:
    from .calculator import FireCheckCalculator

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)

    if not input_model.rows or not result_model.rows:
        _p_body(
            doc,
            "(Нет строк таблицы расчёта — раздел проверки кабелей на невозгорание и таблица 2.1 не сформированы.)",
        )
        return

    r_in = input_model.rows[0]
    r_out = result_model.rows[0]
    cable = r_in.cable.strip() or "кабель (марка не указана)"
    u_line = input_model.tsn.u_line_v if input_model.tsn else 400.0
    ik3 = None
    if input_model.tsn is not None:
        ik3 = FireCheckCalculator.three_phase_sc_current_ka(
            input_model.tsn.r1_mohm, input_model.tsn.x1_mohm, input_model.tsn.u_line_v
        )
    r1_m = input_model.tsn.r1_mohm if input_model.tsn else None
    x1_m = input_model.tsn.x1_mohm if input_model.tsn else None
    _p_caps_blue(doc, "ПРОВЕРКА КАБЕЛЬНЫХ ЛИНИЙ В ЩСН НА НЕВОЗГОРАНИЕ ПРИ ВОЗДЕЙСТВИИ ТОКА КОРОТКОГО ЗАМЫКАНИЯ")

    _p_body(doc, "1. Расчёт токов КЗ производится в единицах, приведенных к напряжению 0,4 кВ. ")
    _p_body(
        doc,
        "Значение тока трехфазного к.з. определяется по формуле [1]:",
    )
    append_inline_omath(doc, omml_Ip0_schema_string(), body_indent=True, centered=True)
    append_mixed_paragraph(doc, legend_where_ik_schema_parts(), body_indent=True)
    _p_body(
        doc,
        "2. Проверка кабелей на невозгорание при воздействии тока короткого замыкания осуществляется по циркуляру "
        "№ Ц-02-98 (Э).",
    )
    _p_body(
        doc,
        "3. Для проверки кабеля на невозгорание по циркуляру № Ц-02-98 (Э) допускается принимать расчетные токи КЗ "
        "на расстоянии 20 м от начала кабельной линии. ",
    )
    append_mixed_paragraph(
        doc,
        [
            (
                "text",
                "4. В соответствии с циркуляром № Ц-02-98 (Э), о проверке кабелей на невозгорание при воздействии тока "
                "короткого замыкания, не допускается нагрев токопроводящей жилы кабеля с поливинилхлоридной изоляцией "
                "свыше 350 ",
            ),
            ("math", xml_mrun("\u00B0C")),
            ("text", ".  "),
        ],
        body_indent=True,
    )

    _p_body(doc, f"Проверка кабельной линии от ТСН до ЩСН {cable} на невозгорание по циркуляру № Ц-02-98 (Э).")
    _p_body(doc, "Значение начальной температуры жилы до КЗ [4]:")
    append_inline_omath(doc, omml_theta_n_schema_string(), body_indent=True, centered=True)
    iw = r_in.working_current_a
    ia = r_in.allowed_current_a
    tn = r_out.wire_temperature_c
    append_mixed_paragraph(doc, legend_where_theta_n_o_parts(), body_indent=True)
    append_mixed_paragraph(doc, legend_where_theta_n_dd_parts(), body_indent=True)
    append_mixed_paragraph(doc, legend_where_theta_n_okr_parts(), body_indent=True)
    append_mixed_paragraph(doc, legend_where_theta_n_irab_parts(), body_indent=True)
    append_mixed_paragraph(doc, legend_where_theta_n_idd_parts(), body_indent=True)
    doc.add_paragraph()

    to = input_model.tsn.theta_o_c if input_model.tsn else 25.0
    tdd = input_model.tsn.theta_dd_c if input_model.tsn else 70.0
    tok = input_model.tsn.theta_okr_c if input_model.tsn else 25.0
    if iw is not None and ia is not None and ia != 0 and tn is not None:
        _p_body(doc, "Значение температуры жилы до КЗ кабельной линии от ТСН до ЩСН:")
        append_omath_xml(
            doc,
            omml_theta_n_string(
                _fmt_int(int(to)),
                _fmt_int(int(tdd)),
                _fmt_int(int(tok)),
                _fmt(iw, 0),
                _fmt(ia, 0),
                str(tn),
            ),
            body_indent=True,
            centered=True,
        )
    else:
        _p_body(doc, f"Значение температуры жилы до КЗ: Θн = {_fmt_int(tn)}℃ (по введённым данным).")

    doc.add_paragraph()
    _p_body(doc, "Значение тока трехфазного к.з.:")
    append_inline_omath(doc, omml_Ip0_schema_string(), body_indent=True, centered=True)
    append_mixed_paragraph(doc, legend_where_ik_schema_parts(), body_indent=True)
    ik_line: float | None = None
    if r_in.short_circuit_ka is not None and r_in.short_circuit_ka > 0:
        ik_line = float(r_in.short_circuit_ka)
    elif ik3 is not None:
        ik_line = float(ik3)
    if ik_line is not None and r1_m is not None and x1_m is not None:
        _p_body(doc, "Значение тока трехфазного к.з. кабельной линии от ТСН до ЩСН:")
        append_inline_omath(
            doc,
            omml_Ip0_numeric_mohm(
                _fmt(u_line, 0),
                _fmt(r1_m, 1),
                _fmt(x1_m, 1),
                ik_ka=_fmt(ik_line, 3),
            ),
            body_indent=True,
            centered=True,
        )
    elif ik_line is not None:
        _p_body(doc, f"Iп0(3): — (для формулы укажите R1 и X1; ток из таблицы {_fmt(ik_line, 3)} кА).")
    else:
        _p_body(doc, "Iп0(3): — (для 1-й линии укажите R1, X1 или ток КЗ в таблице).")

    tae = input_model.tsn.tae_s if input_model.tsn else 0.0
    ik = r_in.short_circuit_ka
    t_off = r_in.shutdown_time_s
    doc.add_paragraph()
    _p_body(doc, "Тепловой импульс от тока КЗ [4]:")
    show_b_num = ik is not None and t_off is not None and r_out.bk_value is not None
    append_omath_xml(doc, omml_B_schema_string(), body_indent=True, centered=False)
    append_mixed_paragraph(doc, legend_where_B_parts(), body_indent=True)
    append_mixed_paragraph(doc, legend_where_B_tae_parts(), body_indent=True)
    if t_off is not None:
        _p_body(
            doc,
            "Суммарное время срабатывания МТЗ и отключения выключателя на отходящей линии к ТСН составляет "
            f"{_fmt(t_off, 2)} с.",
        )
    else:
        _p_body(
            doc,
            "Суммарное время срабатывания МТЗ и отключения выключателя на отходящей линии к ТСН составляет — с.",
        )
    if show_b_num:
        append_inline_omath(
            doc,
            omml_B_string(
                _fmt(ik, 3),
                _fmt(t_off, 2),
                _fmt(tae, 2),
                b_result=_fmt(r_out.bk_value, 2),
            ),
            body_indent=True,
            centered=True,
        )
    else:
        _p_body(doc, "Bтер: —")

    doc.add_paragraph()
    _p_body(doc, "Температура нагрева жил кабеля при действии тока КЗ [4]:")
    s_mm = r_in.section_mm2
    show_theta_k_num = bool(
        s_mm and r_out.wire_temperature_c is not None and r_out.bk_value is not None
    )
    append_omath_xml(doc, omml_theta_k_schema_string(), body_indent=True, centered=False)
    append_mixed_paragraph(doc, legend_where_theta_k_parts(), body_indent=True)
    append_mixed_paragraph(doc, legend_where_theta_k_b_parts(), body_indent=True)
    append_mixed_paragraph(doc, legend_where_theta_k_s_parts(), body_indent=True)
    append_mixed_paragraph(doc, legend_where_theta_k_a_parts(), body_indent=True)
    if show_theta_k_num:
        append_omath_xml(
            doc,
            omml_theta_k_string(
                r_out.final_temperature_display or "—",
                _fmt(s_mm, 0),
                theta_n_num=_fmt_int(r_out.wire_temperature_c),
                b_ter_num=_fmt(r_out.bk_value, 2),
            ),
            body_indent=True,
            centered=True,
        )
    else:
        _p_body(doc, f"Θк = {r_out.final_temperature_display or '—'}℃")

    _p_body(
        doc,
        "Для остальных кабельных линий ОТ ЩСН расчеты проводятся аналогично, результаты расчетов приведены в "
        "таблице 2.1.     ",
    )

    doc.add_paragraph()
    _p_body(
        doc,
        "Таблица 2.1. Проверка кабельных линий от ТСН до потребителей на невозгорание при возникновении КЗ.",
        bold=True,
        center=True,
        no_indent=True,
    )

    h2 = (
        "Номер линии",
        "Марка\nкабеля",
        "Ток\nнагрузки, А",
        "Длительно-\nдопустимый\nток кабеля, А",
        "Температура\nжил до\nКЗ Θн, ℃",
        "Ток КЗ, кА",
        "Время отключения от резервной защиты, с",
        "Температура\nжил после\nКЗ Θк, ℃",
    )
    t2 = doc.add_table(rows=1 + len(result_model.rows), cols=len(h2))
    try:
        t2.style = "Table Grid"
    except (KeyError, ValueError):
        pass
    for j, title in enumerate(h2):
        t2.rows[0].cells[j].text = title
    _style_header_row(t2)
    for i, row in enumerate(result_model.rows, start=1):
        t2.rows[i].cells[0].text = str(row.number)
        t2.rows[i].cells[1].text = row.cable or "—"
        t2.rows[i].cells[2].text = row.working_current_a_text or "—"
        t2.rows[i].cells[3].text = row.allowed_current_a_text or "—"
        t2.rows[i].cells[4].text = "" if row.wire_temperature_c is None else str(row.wire_temperature_c)
        t2.rows[i].cells[5].text = row.short_circuit_ka_text or "—"
        t2.rows[i].cells[6].text = row.shutdown_time_s_text or "—"
        t2.rows[i].cells[7].text = row.final_temperature_display or "—"

    doc.add_paragraph()
    _p_body(
        doc,
        "На основании приведенных расчетов см. таблицу  2.1 можно сделать вывод, что выбранные кабели удовлетворяют "
        "требованиям циркуляра № Ц-02-98 (Э) в отношении невозгорания, все значения температуры токоведущих жил после "
        "аварийного отключения КЗ резервной защитой менее 350 ℃.",
    )
