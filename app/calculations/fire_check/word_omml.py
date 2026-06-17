"""
Уравнения Word через Office Math ML (parse_xml): дроби, корни, степени, индексы.
Шрифт: Cambria Math в блоке формулы.
"""

from __future__ import annotations

import re
import xml.sax.saxutils as xml_esc
from typing import Any, Literal

MixedPart = tuple[Literal["text", "math"], str]

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

MNS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# Как в Raschet_SN.docx: Cambria Math 12 pt (sz 24) внутри m:r, без жирного и курсива.
_MATH_W_RPR = (
    f'<w:rPr xmlns:w="{WNS}">'
    f'<w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math" w:cs="Arial"/>'
    f'<w:b w:val="0"/>'
    f'<w:i w:val="0"/>'
    f'<w:sz w:val="24"/><w:szCs w:val="24"/>'
    f"</w:rPr>"
)

# Знак умножения в формулах Word — как в отчёте СОПТ (_xml_mrun('·')).
MUL_SIGN = "\u00b7"


def normalize_mul_sign(text: str) -> str:
    return (
        text.replace("*", MUL_SIGN)
        .replace("\u2022", MUL_SIGN)
        .replace("\u2219", MUL_SIGN)
    )


def _xe(s: str) -> str:
    return xml_esc.escape(s)


_MATH_MINUS = "\u2212"


def _math_plain_text(s: str) -> str:
    """Тире → математический минус U+2212; «e-(» не трогаем."""
    s = (
        s.replace("\u2103", "\u00B0C")
        .replace("℃", "\u00B0C")
        .replace("\u2014", _MATH_MINUS)
        .replace("\u2013", _MATH_MINUS)
    )
    s = re.sub(r"(?<![eе]) - ", f" {_MATH_MINUS} ", s)
    s = s.replace(f"{_MATH_MINUS}{MUL_SIGN}", MUL_SIGN)
    s = s.replace(f"-{MUL_SIGN}", MUL_SIGN)
    return s


def _wrap_omath(inner: str, *, math_para_center: bool = False) -> str:
    """Обёртка m:oMathPara — в WordprocessingML этот узел должен быть прямым потомком w:p, не w:r."""
    jc = ""
    if math_para_center:
        jc = '<m:oMathParaPr><m:jc m:val="center"/></m:oMathParaPr>'
    return (
        f'<m:oMathPara xmlns:m="{MNS}">'
        f"{jc}"
        f"<m:oMath>{inner}</m:oMath>"
        f"</m:oMathPara>"
    )


def append_omath_xml(
    doc: Any,
    inner_omath: str,
    *,
    centered: bool = True,
    font_pt: int = 14,
    body_indent: bool = False,
) -> None:
    """Добавляет абзац с m:oMathPara (числовые формулы Θн, B, Θк в эталоне)."""
    p = doc.add_paragraph()
    _apply_body_paragraph_format(p, body_indent=body_indent, centered=centered)
    _clear_paragraph_runs(p)

    # Обёртка <p xmlns:m> — иначе parse_xml(m:oMathPara) даёт OMML, который Word не открывает.
    wrapped = _wrap_omath(inner_omath, math_para_center=centered and not body_indent)
    omp = parse_xml(f'<p xmlns:m="{MNS}">{wrapped}</p>')[0]
    p._p.append(omp)


def _apply_body_paragraph_format(
    p: Any,
    *,
    body_indent: bool,
    centered: bool,
) -> None:
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    if body_indent:
        pf.first_line_indent = Cm(1.25)
        pf.left_indent = Cm(0)
        pf.right_indent = Cm(0)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        pf.first_line_indent = Cm(0)
        pf.left_indent = Cm(0)
        pf.right_indent = Cm(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.JUSTIFY


def _clear_paragraph_runs(p: Any) -> None:
    for child in list(p._p):
        if child.tag == qn("w:r"):
            p._p.remove(child)


def _append_text_run(p: Any, text: str) -> None:
    r = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_pr.append(fonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "28")
    r_pr.append(sz)
    sz_cs = OxmlElement("w:szCs")
    sz_cs.set(qn("w:val"), "28")
    r_pr.append(sz_cs)
    r.append(r_pr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    p._p.append(r)


def append_inline_omath(
    doc: Any,
    inner_omath: str,
    *,
    body_indent: bool = True,
    centered: bool = True,
) -> None:
    """Одна формула как m:oMath в абзаце (как схемы Iп0 и Θн в эталоне)."""
    p = doc.add_paragraph()
    _apply_body_paragraph_format(p, body_indent=body_indent, centered=centered)
    _clear_paragraph_runs(p)
    omath = parse_xml(f'<m:oMath xmlns:m="{MNS}">{inner_omath}</m:oMath>')
    p._p.append(omath)


def append_mixed_paragraph(
    doc: Any,
    parts: list[MixedPart],
    *,
    body_indent: bool = True,
    centered: bool = False,
) -> None:
    """Абзац: Times New Roman + встроенные m:oMath (легенды «где…»)."""
    p = doc.add_paragraph()
    _apply_body_paragraph_format(p, body_indent=body_indent, centered=centered)
    _clear_paragraph_runs(p)
    for kind, content in parts:
        if kind == "text":
            if content:
                _append_text_run(p, content)
        elif kind == "math":
            if content:
                omath = parse_xml(f'<m:oMath xmlns:m="{MNS}">{content}</m:oMath>')
                p._p.append(omath)


def xml_mrun(text: str) -> str:
    if text is None:
        text = ""
    safe = _math_plain_text(str(text))
    # Пробелы между токенами формулы сохраняем; «-» только для пустого m:r.
    if safe == "":
        safe = "-"
    t = _xe(safe)
    return (
        f"<m:r>{_MATH_W_RPR}<m:rPr><m:nor/></m:rPr>"
        f"<m:t xml:space=\"preserve\">{t}</m:t></m:r>"
    )


def _omml_sup2(base: str) -> str:
    return xml_msup_expr(xml_mrun(base), "2")


def _omml_sub_var(token: str) -> str:
    """Первая буква — основание, остальное — подстрочный индекс."""
    return xml_msub(token[0], token[1:])


# Длинные обозначения — раньше коротких (Вк.п. до Вк).
_SUBSCRIPT_VARIABLES: tuple[str, ...] = tuple(
    sorted(
        (
            "Вк.п.",
            "Вк.а.",
            "tоткл.",
            "tотк.",
            "Та.эк.",
            "Iп.с.",
            "Iкз(3)",
            "Iраб.",
            "Iраб",
            "Iдд",
            "Iн",
            "Qдд",
            "Qос",
            "Qн",
            "Q0",
            "Qк",
            "Uн",
            "Втер.",
            "Втер",
            "Вк",
            "Smin",
        ),
        key=len,
        reverse=True,
    )
)


def _subscript_var_boundary_ok(text: str, end: int) -> bool:
    if end >= len(text):
        return True
    return text[end] in " =·+-,)(/<>;:" or text[end].isdigit()


def _is_operand_tail_char(ch: str) -> bool:
    if ch.isalnum() or ch in ",.√^":
        return True
    return "\u0400" <= ch <= "\u04ff"


def _operand_start_before_slash(text: str, slash_i: int) -> int:
    j = slash_i - 1
    while j >= 0 and text[j] == " ":
        j -= 1
    if j < 0:
        return 0
    if text[j] == ")":
        depth = 1
        j -= 1
        while j >= 0:
            if text[j] == ")":
                depth += 1
            elif text[j] == "(":
                depth -= 1
                if depth == 0:
                    return j
            j -= 1
        return 0
    while j >= 0 and _is_operand_tail_char(text[j]):
        j -= 1
    return j + 1


def _operand_end_after_slash(text: str, start: int) -> int:
    n = len(text)
    i = start
    while i < n and text[i] == " ":
        i += 1
    if i >= n:
        return n
    if text[i] == "(":
        depth = 1
        j = i + 1
        while j < n and depth > 0:
            if text[j] == "(":
                depth += 1
            elif text[j] == ")":
                depth -= 1
            j += 1
        return j
    depth = 0
    j = i
    while j < n:
        ch = text[j]
        if ch == "(":
            depth += 1
        elif ch == ")":
            if depth == 0:
                break
            depth -= 1
        elif depth == 0 and ch in "/=+":
            break
        elif depth == 0 and ch == "−" and j > i:
            break
        j += 1
    return j


def _find_next_division_slash(text: str, start: int) -> int:
    for i in range(start, len(text)):
        if text[i] != "/":
            continue
        num_start = _operand_start_before_slash(text, i)
        den_end = _operand_end_after_slash(text, i + 1)
        if num_start < i and den_end > i + 1:
            return i
    return -1


def _mm2_per_ka2s_inline_span(s: str, slash: int) -> tuple[int, int] | None:
    """Размерность мм2/кА2·с — в одну строку через /, не стековая дробь."""
    if slash >= 3 and s[slash - 3 : slash] == "мм2":
        mm_start = slash - 3
    elif slash >= 4 and s[slash - 4 : slash] == "мм^2":
        mm_start = slash - 4
    else:
        return None
    den_start = slash + 1
    for den, den_len in (
        ("кА2·с", 5),
        ("кА2с", 4),
        ("кА^2·с", 6),
        ("кА^2с", 5),
    ):
        if s.startswith(den, den_start):
            return mm_start, den_start + den_len
    return None


def _omml_mm2_per_ka2s_inline() -> str:
    return (
        _omml_sup2("мм")
        + xml_mrun("/")
        + _omml_sup2("кА")
        + xml_mrun(MUL_SIGN)
        + xml_mrun("с")
    )


def _parens_balanced(inner: str) -> bool:
    depth = 0
    for ch in inner:
        if ch == "(":
            depth += 1
        elif ch == ")":
            depth -= 1
            if depth < 0:
                return False
    return depth == 0


def _omml_operand_fragment(segment: str) -> str:
    segment = segment.strip()
    if (
        len(segment) >= 2
        and segment[0] == "("
        and segment[-1] == ")"
        and _parens_balanced(segment[1:-1])
    ):
        return _omml_runs_plain_core(segment[1:-1])
    return _omml_runs_plain_core(segment)


def _at_e_exp_start(text: str, i: int) -> bool:
    if i >= len(text) or text[i] != "e":
        return False
    if i + 2 < len(text) and text[i + 1] == "-" and text[i + 2] == "(":
        return True
    return i + 3 < len(text) and text[i + 1] == "^" and text[i + 2] == "-" and text[i + 3] == "("


def _e_exp_close_index(text: str, i: int) -> int:
    """Индекс ')' закрывающей e-(…) или e^-(…); иначе −1."""
    if not _at_e_exp_start(text, i):
        return -1
    if text[i + 1] == "^":
        open_paren = i + 3
    else:
        open_paren = i + 2
    depth = 1
    j = open_paren + 1
    n = len(text)
    while j < n and depth > 0:
        if text[j] == "(":
            depth += 1
        elif text[j] == ")":
            depth -= 1
        j += 1
    if depth != 0:
        return -1
    return j - 1


def _find_next_e_exp_start(text: str, start: int) -> int:
    for k in range(start, len(text)):
        if _at_e_exp_start(text, k):
            return k
    return -1


def _omml_core_scan(s: str) -> str:
    """Сначала e-(…) / e^-(…) как надстрочная степень, затем дроби a/b."""
    if not s:
        return ""
    out: list[str] = []
    i = 0
    n = len(s)
    while i < n:
        if _at_e_exp_start(s, i):
            close = _e_exp_close_index(s, i)
            if close < 0:
                out.append(_omml_with_stacked_fractions(s[i]))
                i += 1
                continue
            if s[i + 1] == "^":
                power = s[i + 2 : close + 1]
            else:
                power = s[i + 1 : close + 1]
            out.append(xml_mexp(_omml_with_stacked_fractions(power)))
            i = close + 1
            continue
        next_e = _find_next_e_exp_start(s, i + 1)
        chunk_end = next_e if next_e >= 0 else n
        out.append(_omml_with_stacked_fractions(s[i:chunk_end]))
        i = chunk_end
    return "".join(out)


def _omml_with_stacked_fractions(s: str) -> str:
    """Косая черта / → стековая дробь m:f; мм2/кА2·с — исключение, в одну строку."""
    if not s:
        return ""
    slash = _find_next_division_slash(s, 0)
    if slash < 0:
        return _omml_runs_plain_tokens(s)
    unit_span = _mm2_per_ka2s_inline_span(s, slash)
    if unit_span is not None:
        mm_start, unit_end = unit_span
        before = _omml_runs_plain_tokens(s[:mm_start]) if mm_start > 0 else ""
        after = _omml_with_stacked_fractions(s[unit_end:])
        return before + _omml_mm2_per_ka2s_inline() + after
    num_start = _operand_start_before_slash(s, slash)
    den_start = slash + 1
    while den_start < len(s) and s[den_start] == " ":
        den_start += 1
    den_end = _operand_end_after_slash(s, den_start)
    before = _omml_runs_plain_tokens(s[:num_start]) if num_start > 0 else ""
    num = _omml_operand_fragment(s[num_start:slash])
    den = _omml_operand_fragment(s[den_start:den_end])
    after = _omml_with_stacked_fractions(s[den_end:])
    return before + xml_mfrac(num, den) + after


def _split_squared_paren_groups(text: str) -> list[tuple[str, str]]:
    """Сегменты: ('plain', …) и ('sq', внутренность) для выражений (….)2."""
    segments: list[tuple[str, str]] = []
    plain_start = 0
    i = 0
    n = len(text)
    while i < n:
        if text[i] != "(":
            i += 1
            continue
        depth = 1
        j = i + 1
        while j < n and depth > 0:
            if text[j] == "(":
                depth += 1
            elif text[j] == ")":
                depth -= 1
            j += 1
        if depth != 0:
            break
        close = j - 1
        if j < n and text[j] == "^" and j + 1 < n and text[j + 1] == "2":
            next_ok = j + 2 >= n or text[j + 2] in " ·=,;)"
            if next_ok:
                if plain_start < i:
                    segments.append(("plain", text[plain_start:i]))
                segments.append(("sq", text[i + 1 : close]))
                i = j + 2
                plain_start = i
                continue
        if j < n and text[j] == "2":
            next_ok = j + 1 >= n or text[j + 1] in " ·=,;)"
            if next_ok:
                if plain_start < i:
                    segments.append(("plain", text[plain_start:i]))
                segments.append(("sq", text[i + 1 : close]))
                i = j + 1
                plain_start = i
                continue
        i += 1
    if plain_start < n:
        segments.append(("plain", text[plain_start:n]))
    return segments


def _omml_runs_plain_tokens(s: str) -> str:
    """Подстрочные индексы, мм², кА²с и прочие степени 2 (без (….)2 и без /)."""
    if not s:
        return ""
    out: list[str] = []
    buf: list[str] = []
    i = 0
    n = len(s)

    def flush_buf() -> None:
        if buf:
            out.append(xml_mrun("".join(buf)))
            buf.clear()

    def try_subscript_variable() -> bool:
        nonlocal i
        for token in _SUBSCRIPT_VARIABLES:
            if not s.startswith(token, i):
                continue
            end = i + len(token)
            if token == "Iп.с." and end < n and s[end] in "2^":
                continue
            if not _subscript_var_boundary_ok(s, end):
                continue
            flush_buf()
            out.append(_omml_sub_var(token))
            i = end
            return True
        return False

    def try_e_exponent() -> bool:
        """e-((2·tотк.)/(Та.эк.)) — показатель степени надстрочный."""
        nonlocal i
        if i >= n or s[i] != "e" or i + 2 >= n or s[i + 1] != "-" or s[i + 2] != "(":
            return False
        depth = 1
        j = i + 3
        while j < n and depth > 0:
            if s[j] == "(":
                depth += 1
            elif s[j] == ")":
                depth -= 1
            j += 1
        if depth != 0:
            return False
        close = j - 1
        power = s[i + 1 : close + 1]
        flush_buf()
        out.append(xml_mexp(_omml_runs_plain_core(power)))
        i = close + 1
        return True

    def try_cyrillic_e_power() -> bool:
        """е^k, е^0,07 — показатель степени надстрочный."""
        nonlocal i
        if i >= n or s[i] != "е":
            return False
        if i + 1 < n and s[i + 1] == "^":
            m = re.match(r"\^([k\d,]+)", s[i + 1 :])
            if m:
                flush_buf()
                out.append(xml_msup_expr(xml_mrun("е"), m.group(1)))
                i += 1 + len(m.group(0))
                return True
        if i + 1 < n and s[i + 1] == "k" and (i + 2 >= n or s[i + 2] in " ·+=,;)(−-"):
            flush_buf()
            out.append(xml_msup_expr(xml_mrun("е"), "k"))
            i += 2
            return True
        m = re.match(r"е([\d,]+)", s[i:])
        if m and (i + len(m.group(0)) >= n or s[i + len(m.group(0))] in " ·+=,;)(−-"):
            flush_buf()
            out.append(xml_msup_expr(xml_mrun("е"), m.group(1)))
            i += len(m.group(0))
            return True
        return False

    while i < n:
        if s.startswith("Iп.с.^2", i) and (i + 7 >= n or s[i + 7] in " ·=,;)"):
            flush_buf()
            out.append(xml_msup_expr(xml_msub("I", "п.с."), "2"))
            i += 7
            continue
        if s.startswith("Iп.с.2", i) and (i + 6 >= n or s[i + 6] in " ·=,;)"):
            flush_buf()
            out.append(xml_msup_expr(xml_msub("I", "п.с."), "2"))
            i += 6
            continue
        if try_e_exponent():
            continue
        if try_cyrillic_e_power():
            continue
        if s.startswith("мм^2", i):
            flush_buf()
            out.append(_omml_sup2("мм"))
            i += 4
            continue
        if s.startswith("мм2", i):
            flush_buf()
            out.append(_omml_sup2("мм"))
            i += 3
            continue
        if s.startswith("кА^2с", i):
            flush_buf()
            out.append(_omml_sup2("кА"))
            out.append(xml_mrun("с"))
            i += 5
            continue
        if (
            s.startswith("кА^2", i)
            and i + 5 < n
            and s[i + 4] == MUL_SIGN
            and s[i + 5] == "с"
        ):
            flush_buf()
            out.append(_omml_sup2("кА"))
            out.append(xml_mrun(MUL_SIGN))
            out.append(xml_mrun("с"))
            i += 6
            continue
        if s.startswith("кА2с", i):
            flush_buf()
            out.append(_omml_sup2("кА"))
            out.append(xml_mrun("с"))
            i += 4
            continue
        if (
            s.startswith("кА2", i)
            and i + 4 < n
            and s[i + 3] == MUL_SIGN
            and s[i + 4] == "с"
        ):
            flush_buf()
            out.append(_omml_sup2("кА"))
            out.append(xml_mrun(MUL_SIGN))
            out.append(xml_mrun("с"))
            i += 5
            continue
        if s.startswith("S^2", i) and (i + 3 >= n or s[i + 3] in " =/,)"):
            flush_buf()
            out.append(_omml_sup2("S"))
            i += 3
            continue
        if s.startswith("S2", i) and (i + 2 >= n or s[i + 2] in " =/,)"):
            flush_buf()
            out.append(_omml_sup2("S"))
            i += 2
            continue
        caret_sq = re.match(r"([\d,]+)\^2(?=[ ·=,;)]|$)", s[i:])
        if caret_sq:
            flush_buf()
            out.append(_omml_sup2(caret_sq.group(1)))
            i += len(caret_sq.group(0))
            continue
        if try_subscript_variable():
            continue
        if s[i] == MUL_SIGN:
            flush_buf()
            out.append(xml_mrun(MUL_SIGN))
            i += 1
            continue
        buf.append(s[i])
        i += 1

    flush_buf()
    return "".join(out)


def _omml_runs_plain_core(s: str) -> str:
    return _omml_core_scan(s)


def omml_runs_from_plain_text(text: str) -> str:
    """Фрагмент формулы: подстрочные индексы, (….)², мм², кА²с."""
    if not text:
        return ""
    s = _math_plain_text(normalize_mul_sign(text)).replace("мм²", "мм2").replace("кА²", "кА2")
    chunks: list[str] = []
    for kind, content in _split_squared_paren_groups(s):
        if kind == "plain":
            if content:
                chunks.append(_omml_runs_plain_core(content))
        else:
            inner = _omml_runs_plain_core(content)
            chunks.append(xml_msup_expr(xml_mdelim(inner), "2"))
    return "".join(chunks)


def omml_plain_formula(text: str) -> str:
    """Плоский текст формулы → OMML (подстрочные/надстрочные, · отдельным m:r)."""
    return omml_runs_from_plain_text(normalize_mul_sign(text))


def omml_plain_formula_inline(text: str) -> str:
    """Inline-формула в одном абзаце с текстом — тот же размер, что и display."""
    return omml_plain_formula(text)


def xml_msubsup(base: str, sub: str, sup: str) -> str:
    """Iп0⁽³⁾ — m:sSubSup как в Raschet_SN.docx."""
    return (
        "<m:sSubSup>"
        f"<m:e>{xml_mrun(base)}</m:e>"
        f"<m:sub>{xml_mrun(sub)}</m:sub>"
        f"<m:sup>{xml_mrun(sup)}</m:sup>"
        "</m:sSubSup>"
    )


def xml_mdelim(inner: str) -> str:
    """Круглые скобки m:d — как (ΘДД−ΘОКР) в эталоне."""
    return f"<m:d><m:dPr/><m:e>{inner}</m:e></m:d>"


def xml_msub(base: str, sub: str) -> str:
    """Подстрочный индекс: корень должен быть m:sSub (а не m:sub), иначе Word не рисует формулу."""
    return (
        "<m:sSub>"
        f"<m:e>{xml_mrun(base)}</m:e>"
        f"<m:sub>{xml_mrun(sub)}</m:sub>"
        "</m:sSub>"
    )


def xml_msup(base: str, sup: str) -> str:
    # m:sup — как m:sub: без лишнего m:e, иначе Word не открывает docx.
    return (
        "<m:sSup>"
        f"<m:e>{xml_mrun(base)}</m:e>"
        f"<m:sup>{xml_mrun(sup)}</m:sup>"
        "</m:sSup>"
    )


def xml_msup_expr(base_xml: str, sup: str) -> str:
    """Степень, когда основание — уже фрагмент OMML."""
    return (
        "<m:sSup>"
        f"<m:e>{base_xml}</m:e>"
        f"<m:sup>{xml_mrun(sup)}</m:sup>"
        "</m:sSup>"
    )


def xml_mfrac(num_inner: str, den_inner: str) -> str:
    return (
        "<m:f>"
        f"<m:num><m:e>{num_inner}</m:e></m:num>"
        f"<m:den><m:e>{den_inner}</m:e></m:den>"
        "</m:f>"
    )


def xml_msqrt(inner: str) -> str:
    """Корень без видимого индекса степени (degHide) — иначе в Word пустые квадраты."""
    return (
        "<m:rad>"
        "<m:radPr><m:degHide m:val=\"1\"/></m:radPr>"
        "<m:deg/>"
        f"<m:e>{inner}</m:e>"
        "</m:rad>"
    )


def xml_mexp(power_inner: str) -> str:
    # m:sup — без лишнего m:e (как xml_msup), иначе Word не открывает docx.
    return (
        "<m:sSup>"
        f"<m:e>{xml_mrun('e')}</m:e>"
        f"<m:sup>{power_inner}</m:sup>"
        "</m:sSup>"
    )


def omml_B_string(
    i_ka: str,
    t_s: str,
    tae_s: str,
    *,
    b_result: str | None = None,
) -> str:
    """Bтер = I²∙t + Tаэ; при b_result — итог и размерность в конце: … = 27,55 кА²·с."""
    b = xml_msub("B", "тер")
    i2 = xml_msup_expr(xml_mrun(i_ka), "2")
    dot = xml_mrun("\u2219")
    s = (
        f"{b}{xml_mrun(' = ')}{i2}{dot}{xml_mrun(t_s)}"
        f"{xml_mrun(' + ')}{xml_mrun(tae_s)}"
    )
    if b_result is not None:
        ka2 = xml_msup_expr(xml_mrun("кА"), "2")
        s += f"{xml_mrun(' = ')}{xml_mrun(b_result)}{ka2}{xml_mrun('\u00b7')}{xml_mrun('с')}"
    return s


def omml_B_schema_string() -> str:
    """Bтер = Iп0² · tоткл + Tаэ (без чисел)."""
    b = xml_msub("B", "тер")
    ip0 = xml_msub("I", "п0")
    i2 = xml_msup_expr(ip0, "2")
    toff = xml_msub("t", "откл")
    ta = xml_msub("T", "аэ")
    return f"{b}{xml_mrun(' = ')}{i2}{xml_mrun('·')}{toff}{xml_mrun(' + ')}{ta}"


def omml_theta_n_string(to_c: str, tdd_c: str, tok_c: str, iw: str, ia: str, tn: str) -> str:
    th = xml_msub("Θ", "н")
    frac = xml_mfrac(xml_mrun(iw), xml_mrun(ia))
    sq = xml_msup_expr(xml_mdelim(frac), "2")
    diff = f"{xml_mrun(tdd_c)}{xml_mrun('−')}{xml_mrun(tok_c)}"
    return (
        f"{th}{xml_mrun(' = ')}"
        f"{xml_mrun(to_c)}{xml_mrun('+')}{xml_mdelim(diff)}"
        f"{xml_mrun('·')}{sq}{xml_mrun(' = ')}{xml_mrun(tn)}{xml_mrun(' \u00B0C')}"
    )


def omml_Ip0_schema_string() -> str:
    """Iп0⁽³⁾ = U_ном / (√3 · √(R₁эк²+X₁эк²)) — как в тексте до подстановки чисел."""
    num = xml_msub("U", "ном")
    rsq = xml_msup_expr(xml_msub("R", "1эк"), "2")
    xsq = xml_msup_expr(xml_msub("X", "1эк"), "2")
    den = f"{xml_msqrt(xml_mrun('3'))}{xml_mrun('·')}{xml_msqrt(rsq + xml_mrun('+') + xsq)}"
    ip = xml_msubsup("I", "п0", "(3)")
    return f"{ip}{xml_mrun(' = ')}{xml_mfrac(num, den)}"


def omml_Ip0_numeric_mohm(
    u_nom_v: str,
    r1_mohm: str,
    x1_mohm: str,
    *,
    ik_ka: str | None = None,
) -> str:
    """Iп0⁽³⁾ = Uном / (√3·√(R₁²+X₁²)); при ik_ka — итог в конце: … = 5,247 кА."""
    ip = xml_msubsup("I", "п0", "(3)")
    rsq = xml_msup_expr(xml_mrun(r1_mohm), "2")
    xsq = xml_msup_expr(xml_mrun(x1_mohm), "2")
    den = f"{xml_msqrt(xml_mrun('3'))}{xml_mrun('\u2219')}{xml_msqrt(rsq + xml_mrun('+') + xsq)}"
    s = f"{ip}{xml_mrun(' = ')}{xml_mfrac(xml_mrun(u_nom_v), den)}"
    if ik_ka is not None:
        s += f"{xml_mrun(' = ')}{xml_mrun(ik_ka)}{xml_mrun(' кА')}"
    return s


def _omml_theta_k_exp_pair(b_inner: str, s_mm: str) -> tuple[str, str]:
    """e^(19,58·B/S²) и e^(19,58·B/S²−1) — стековая дробь, как в численной Θк."""
    s2 = xml_msup_expr(xml_mrun(s_mm), "2")
    num = f"{xml_mrun('19,58')}{xml_mrun('·')}{b_inner}"
    frac = xml_mfrac(num, s2)
    pow2 = f"{frac}{xml_mrun('−1')}"
    return xml_mexp(frac), xml_mexp(pow2)


def omml_theta_k_schema_string(s_mm: str = "S") -> str:
    """Θк = Θн·e^κ + a·(e^κ−1), κ = 19,58·Bтер/S²."""
    thk = xml_msub("Θ", "к")
    thn = xml_msub("Θ", "н")
    e1, _ = _omml_theta_k_exp_pair(xml_msub("B", "тер"), s_mm)
    return (
        f"{thk}{xml_mrun(' = ')}{thn}{xml_mrun('·')}{e1}"
        f"{xml_mrun(' + ')}{xml_mrun('a')}{xml_mrun('·(')}{e1}{xml_mrun(' − 1)')}"
    )


def omml_theta_k_string(
    theta_k: str,
    s_mm: str,
    *,
    theta_n_num: str | None = None,
    b_ter_num: str | None = None,
) -> str:
    """Θк; при theta_n_num и b_ter_num — числа в показателях и перед e, иначе символы Θн и Bтер."""
    thk = xml_msub("Θ", "к")
    lead = xml_mrun(theta_n_num) if theta_n_num else xml_msub("Θ", "н")
    b_piece = xml_mrun(b_ter_num) if b_ter_num else xml_msub("B", "тер")
    e1, _ = _omml_theta_k_exp_pair(b_piece, s_mm)
    return (
        f"{thk}{xml_mrun(' = ')}{lead}{xml_mrun('·')}{e1}{xml_mrun(' + 228·(')}{e1}{xml_mrun(' − 1)')}"
        f"{xml_mrun(' = ')}{xml_mrun(theta_k)}{xml_mrun(' \u00B0C')}"
    )


def omml_theta_n_schema_string() -> str:
    """Θн = Θо + (ΘДД − ΘОКР)·(Iраб/IДД)²."""
    thn = xml_msub("Θ", "н")
    tho = xml_msub("Θ", "о")
    tdd = xml_msub("Θ", "ДД")
    tok = xml_msub("Θ", "ОКР")
    irab = xml_msub("I", "раб")
    idd = xml_msub("I", "ДД")
    frac = xml_mfrac(irab, idd)
    sq = xml_msup_expr(xml_mdelim(frac), "2")
    diff = f"{tdd}{xml_mrun('−')}{tok}"
    return f"{thn}{xml_mrun(' = ')}{tho}{xml_mrun('+')}{xml_mdelim(diff)}{xml_mrun('·')}{sq}{xml_mrun(';')}"


def legend_where_ik_schema_parts() -> list[MixedPart]:
    """где Uном … R1эк и X1эк … (как в Raschet_SN)."""
    return [
        ("text", "где "),
        ("math", xml_msub("U", "ном")),
        ("text", "- номинальное напряжение сети (фазное значение); "),
        ("math", xml_msub("R", "1эк")),
        ("text", " и "),
        ("math", xml_msub("X", "1эк")),
        (
            "text",
            " - соответственно эквивалентное активное и индуктивное сопротивление "
            "прямой последовательности до точки к.з.",
        ),
    ]


def legend_where_theta_n_o_parts() -> list[MixedPart]:
    return [
        ("text", "где  "),
        ("math", xml_msub("Θ", "о")),
        ("text", " - температура окружающей среды во время КЗ,  "),
        ("math", xml_msub("Θ", "о")),
        ("text", " принимаем равной 25"),
        ("math", xml_mrun("\u00B0C")),
        ("text", ";"),
    ]


def legend_where_theta_n_dd_parts() -> list[MixedPart]:
    return [
        ("math", xml_msub("Θ", "ДД")),
        (
            "text",
            " – значение расчетной длительно допустимой температуры жилы, для кабелей "
            "с поливинилхлоридной изоляцией по циркулярам № Ц-02-98 (Э) принимаем равной 70 ",
        ),
        ("math", xml_mrun("\u00B0C")),
        ("text", ";"),
    ]


def legend_where_theta_n_okr_parts() -> list[MixedPart]:
    return [
        ("math", xml_msub("Θ", "ОКР")),
        ("text", " – значение расчетной температуры окружающей среды 25 "),
        ("math", xml_mrun("\u00B0C")),
        ("text", ";"),
    ]


def legend_where_theta_n_irab_parts() -> list[MixedPart]:
    return [
        ("math", xml_msub("I", "раб")),
        ("text", " – значение тока до КЗ, А(выбран по макс. нагрузке на ТСН);"),
    ]


def legend_where_theta_n_idd_parts() -> list[MixedPart]:
    return [
        ("math", xml_msub("I", "ДД")),
        ("text", " – значение расчетного длительно допустимого тока в соответствии с ПУЭ, А."),
    ]


def legend_where_B_parts() -> list[MixedPart]:
    return [
        ("text", "где "),
        ("math", xml_msub("I", "п0")),
        ("text", "– значение трехфазного тока КЗ в конце участка кабельной линии от ТСН до ЩСН, кА; "),
        ("math", xml_msub("t", "откл")),
        ("text", " – время отключения тока КЗ, с;"),
    ]


def legend_where_B_tae_parts() -> list[MixedPart]:
    return [
        ("text", "- эквивалентная постоянная времени затухания апериодической составляющей тока КЗ от удаленных источников "),
        ("text", "по циркуляру № Ц-02-98 (Э) для сети 0,4 кВ принимаем равной "),
        ("math", xml_mrun("0,02")),
        ("text", " с."),
    ]


def legend_where_theta_k_parts() -> list[MixedPart]:
    return [
        ("text", "где "),
        ("math", xml_msub("Θ", "н")),
        ("text", "– температура жилы до КЗ, "),
        ("math", xml_mrun("\u00B0C")),
        ("text", "; в – постоянная, характеризующая теплофизические характеристики материала жилы, для меди 19,58 "),
        ("math", xml_mrun("мм\u2074/(кА\u00B2\u2219с)")),
        ("text", ";"),
    ]


def legend_where_theta_k_b_parts() -> list[MixedPart]:
    return [
        ("text", " "),
        ("math", xml_msub("B", "тер")),
        ("text", " – тепловой импульс от тока КЗ, "),
        ("math", xml_mrun("кА\u00B2\u2219с")),
        ("text", ";"),
    ]


def legend_where_theta_k_s_parts() -> list[MixedPart]:
    return [
        ("text", "S – сечение жилы кабельной линии, "),
        ("math", xml_mrun("мм\u00B2")),
        ("text", ";"),
    ]


def legend_where_theta_k_a_parts() -> list[MixedPart]:
    return [
        ("text", "a – величина, обратная температурному коэффициенту электрического сопротивления при 0 "),
        ("math", xml_mrun("\u00B0C")),
        ("text", ", равная 228 "),
        ("math", xml_mrun("\u00B0C")),
        ("text", ". "),
    ]
