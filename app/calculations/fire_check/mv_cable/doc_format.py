"""Оформление Word-отчёта КЛ свыше 1 кВ — поля, шрифты, стили из шаблона."""

from __future__ import annotations

from pathlib import Path
import re
import shutil
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from app.calculations.fire_check.word_omml import (
    MNS,
    append_mixed_paragraph,
    append_omath_xml,
    omml_plain_formula,
    omml_plain_formula_inline,
)

_PAGE_WIDTH_CM = 21.0
_PAGE_HEIGHT_CM = 29.7
_PAGE_MARGIN_TOP_CM = 1.69
_PAGE_MARGIN_BOTTOM_CM = 2.0
_PAGE_MARGIN_LEFT_CM = 2.25
_PAGE_MARGIN_RIGHT_CM = 1.25
_PAGE_MARGIN_GUTTER_CM = 0.0

_TEXT_FONT_NAME = "Arial"
_FONT_SIZE_BODY_PT = 12
_FONT_SIZE_SECTION_PT = 14
_FONT_SIZE_SUBSECTION_PT = 13

_STYLE_NORMAL = "Normal"
_STYLE_HEADING_1 = "Heading 1"
_STYLE_HEADING_2 = "Heading 2"
_STYLE_SOURCE_LIST = "Список исполнителей и приложения"

_BODY_FIRST_LINE_INDENT_CM = 1.0
_NUMBERED_SECTION_SPACE_BEFORE_PT = 12
SECTION_SPACE_BEFORE_PT = _NUMBERED_SECTION_SPACE_BEFORE_PT
_NUMBERED_SECTION_RE = re.compile(r"^(\d+)\)")
_SOURCE_LIST_NUM_ID = 25

_LEGEND_SEPS = ("  - ", " – ", " - ")

_ASSETS_DIR = Path(__file__).resolve().parent / "assets"
_BUNDLED_TEMPLATE = _ASSETS_DIR / "mv_cable_report_template.docx"
_USER_TEMPLATE = Path(r"C:\Users\Николаев\Documents\Выбор и проверка КЛ 6 кВ до ТСН.docx")


def resolve_template_path() -> Path:
    if _BUNDLED_TEMPLATE.is_file():
        return _BUNDLED_TEMPLATE
    if _USER_TEMPLATE.is_file():
        _ASSETS_DIR.mkdir(parents=True, exist_ok=True)
        shutil.copy2(_USER_TEMPLATE, _BUNDLED_TEMPLATE)
        return _BUNDLED_TEMPLATE
    raise FileNotFoundError(
        f"Шаблон Word не найден: {_BUNDLED_TEMPLATE} или {_USER_TEMPLATE}"
    )


def open_template_document() -> Document:
    doc = Document(str(resolve_template_path()))
    _clear_document_body(doc)
    _clear_headers_footers(doc)
    _apply_page_setup(doc)
    return doc


def _clear_header_footer_part(part: Any) -> None:
    if part is None:
        return
    element = part._element
    for child in list(element):
        element.remove(child)


def _clear_headers_footers(doc: Document) -> None:
    for section in doc.sections:
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        if section.different_first_page_header_footer:
            _clear_header_footer_part(section.first_page_header)
            _clear_header_footer_part(section.first_page_footer)
        _clear_header_footer_part(section.header)
        _clear_header_footer_part(section.footer)


def _clear_document_body(doc: Document) -> None:
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split("}")[-1]
        if tag in ("p", "tbl"):
            body.remove(child)


def _apply_page_setup(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(_PAGE_WIDTH_CM)
    section.page_height = Cm(_PAGE_HEIGHT_CM)
    section.top_margin = Cm(_PAGE_MARGIN_TOP_CM)
    section.bottom_margin = Cm(_PAGE_MARGIN_BOTTOM_CM)
    section.left_margin = Cm(_PAGE_MARGIN_LEFT_CM)
    section.right_margin = Cm(_PAGE_MARGIN_RIGHT_CM)
    section.gutter = Cm(_PAGE_MARGIN_GUTTER_CM)


def _set_run_font(run: Any, *, size_pt: int = _FONT_SIZE_BODY_PT, bold: bool = False) -> None:
    run.font.name = _TEXT_FONT_NAME
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), _TEXT_FONT_NAME)
    r_fonts.set(qn("w:hAnsi"), _TEXT_FONT_NAME)
    r_fonts.set(qn("w:cs"), _TEXT_FONT_NAME)
    r_fonts.set(qn("w:eastAsia"), _TEXT_FONT_NAME)
    b = r_pr.find(qn("w:b"))
    if b is None:
        b = OxmlElement("w:b")
        r_pr.append(b)
    b.set(qn("w:val"), "0" if not bold else "1")


def _apply_paragraph_runs_font(
    p: Any,
    *,
    size_pt: int = _FONT_SIZE_BODY_PT,
    bold: bool = False,
) -> None:
    for run in p.runs:
        _set_run_font(run, size_pt=size_pt, bold=bold)


def _apply_paragraph_numbering(p: Any, *, num_id: int, ilvl: int = 0) -> None:
    p_pr = p._p.get_or_add_pPr()
    old = p_pr.find(qn("w:numPr"))
    if old is not None:
        p_pr.remove(old)
    num_pr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl_el)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)


def _omml_plain(text: str) -> str:
    return omml_plain_formula(text)


def _omml_plain_inline(text: str) -> str:
    return omml_plain_formula_inline(text)


def _clear_paragraph_runs(p: Any) -> None:
    for child in list(p._p):
        if child.tag == qn("w:r"):
            p._p.remove(child)


def _populate_mixed_paragraph(p: Any, parts: list[tuple[str, str]]) -> None:
    _clear_paragraph_runs(p)
    for kind, content in parts:
        if not content:
            continue
        if kind == "text":
            run = p.add_run(content)
            _set_run_font(run, bold=False)
        elif kind == "math":
            omath = parse_xml(
                f'<m:oMath xmlns:m="{MNS}">{_omml_plain_inline(content)}</m:oMath>'
            )
            p._p.append(omath)


def _split_formula_legend(text: str) -> tuple[str, str | None]:
    for sep in _LEGEND_SEPS:
        idx = text.find(sep)
        if idx > 0:
            left = text[:idx]
            if any(ch in left for ch in "=<>"):
                return left, text[idx:]
    return text, None


def _fix_paragraph_spacing(p: Any, *, first_line_indent_cm: float) -> None:
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(first_line_indent_cm)


def _fix_legend_runs_arial(p: Any) -> None:
    for run in p.runs:
        _set_run_font(run, bold=False)


def add_centered_formula(doc: Document, text: str) -> None:
    append_omath_xml(
        doc,
        _omml_plain(text),
        centered=True,
        body_indent=False,
        font_pt=_FONT_SIZE_BODY_PT,
    )
    _fix_paragraph_spacing(doc.paragraphs[-1], first_line_indent_cm=0)


def add_body_formula(
    doc: Document,
    text: str,
    *,
    first_line_indent_cm: float = _BODY_FIRST_LINE_INDENT_CM,
) -> None:
    left, legend = _split_formula_legend(text)
    if legend:
        append_mixed_paragraph(
            doc,
            [("math", _omml_plain_inline(left)), ("text", legend)],
            body_indent=True,
            centered=False,
        )
        _fix_legend_runs_arial(doc.paragraphs[-1])
    else:
        append_omath_xml(
            doc,
            _omml_plain(text),
            centered=False,
            body_indent=True,
            font_pt=_FONT_SIZE_BODY_PT,
        )
    _fix_paragraph_spacing(doc.paragraphs[-1], first_line_indent_cm=first_line_indent_cm)


def add_body_mixed_paragraph(
    doc: Document,
    parts: list[tuple[str, str]],
    *,
    first_line_indent_cm: float = _BODY_FIRST_LINE_INDENT_CM,
) -> None:
    mixed: list[tuple[str, str]] = []
    for kind, content in parts:
        if kind == "math":
            mixed.append(("math", _omml_plain_inline(content)))
        else:
            mixed.append(("text", content))
    append_mixed_paragraph(doc, mixed, body_indent=True, centered=False)
    _fix_legend_runs_arial(doc.paragraphs[-1])
    _fix_paragraph_spacing(doc.paragraphs[-1], first_line_indent_cm=first_line_indent_cm)


def _style_body_paragraph(p: Any, *, first_line_indent_cm: float = _BODY_FIRST_LINE_INDENT_CM) -> None:
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(first_line_indent_cm)
    pf.left_indent = Cm(0)
    pf.right_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_blank_line(doc: Document) -> None:
    p = doc.add_paragraph("", style=_STYLE_NORMAL)
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(0)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15


def add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style=_STYLE_HEADING_1)
    _apply_paragraph_runs_font(p, size_pt=_FONT_SIZE_SECTION_PT, bold=True)
    add_blank_line(doc)


def add_subsection_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style=_STYLE_HEADING_2)
    _apply_paragraph_runs_font(p, size_pt=_FONT_SIZE_SUBSECTION_PT, bold=True)
    add_blank_line(doc)


def add_body_paragraph(
    doc: Document,
    text: str,
    *,
    first_line_indent_cm: float = _BODY_FIRST_LINE_INDENT_CM,
    space_before_pt: float | None = None,
) -> None:
    p = doc.add_paragraph(text, style=_STYLE_NORMAL)
    _style_body_paragraph(p, first_line_indent_cm=first_line_indent_cm)
    if space_before_pt is not None:
        p.paragraph_format.space_before = Pt(space_before_pt)
    else:
        m = _NUMBERED_SECTION_RE.match(text.lstrip())
        if m and int(m.group(1)) > 1:
            p.paragraph_format.space_before = Pt(_NUMBERED_SECTION_SPACE_BEFORE_PT)
    _apply_paragraph_runs_font(p, bold=False)


def add_body_verdict_paragraph(
    doc: Document,
    *,
    passed: bool,
    prefix: str,
    suffix: str,
    first_line_indent_cm: float = _BODY_FIRST_LINE_INDENT_CM,
    space_before_pt: float | None = SECTION_SPACE_BEFORE_PT,
) -> None:
    """Вывод: … соответствует / не соответствует … (при несоответствии фраза жирная)."""
    p = doc.add_paragraph(style=_STYLE_NORMAL)
    _style_body_paragraph(p, first_line_indent_cm=first_line_indent_cm)
    if space_before_pt is not None:
        p.paragraph_format.space_before = Pt(space_before_pt)

    _set_run_font(p.add_run(prefix), bold=False)
    match = "соответствует" if passed else "не соответствует"
    _set_run_font(p.add_run(match), bold=not passed)
    _set_run_font(p.add_run(f" {suffix}"), bold=False)


def add_centered_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style=_STYLE_NORMAL)
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(0)
    pf.right_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _apply_paragraph_runs_font(p, bold=False)


def add_source_data_item(doc: Document, text: str) -> None:
    try:
        p = doc.add_paragraph(style=_STYLE_SOURCE_LIST)
    except KeyError:
        p = doc.add_paragraph(style=_STYLE_NORMAL)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _apply_paragraph_numbering(p, num_id=_SOURCE_LIST_NUM_ID, ilvl=0)
    _populate_mixed_paragraph(p, [("text", text)])


def add_source_data_mixed_item(doc: Document, parts: list[tuple[str, str]]) -> None:
    try:
        p = doc.add_paragraph(style=_STYLE_SOURCE_LIST)
    except KeyError:
        p = doc.add_paragraph(style=_STYLE_NORMAL)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _apply_paragraph_numbering(p, num_id=_SOURCE_LIST_NUM_ID, ilvl=0)
    _populate_mixed_paragraph(p, parts)
